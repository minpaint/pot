import logging
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate

from directory.document_generators.base import prepare_employee_context
from directory.utils import find_appropriate_commission, get_commission_members_formatted
from directory.utils.declension import decline_phrase

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(
    "/home/django/webapps/potby/media/document_templates/etalon/udostoverenie_final.docx"
)
_TEMPLATE_BYTES = None


def _get_template_bytes() -> bytes:
    global _TEMPLATE_BYTES
    if _TEMPLATE_BYTES is None:
        _TEMPLATE_BYTES = TEMPLATE_PATH.read_bytes()
    return _TEMPLATE_BYTES


def _build_certificate_context(employee, commission_cache: Optional[Dict[int, Dict[str, Any]]] = None) -> Dict[str, Any]:
    context = prepare_employee_context(employee)
    commission = find_appropriate_commission(employee)
    cache_key = commission.id if commission else None

    if commission_cache is not None and cache_key in commission_cache:
        cached = commission_cache[cache_key]
        context.setdefault('chairman_name_initials', cached.get('chairman_name_initials', '—'))
        context.setdefault('vice_chairman_name_initials', cached.get('vice_chairman_name_initials', '—'))
        context.setdefault('binding_name_genitive', cached.get('binding_name_genitive', ''))
        return context

    commission_data = get_commission_members_formatted(commission) if commission else {}
    chairman = commission_data.get('chairman', {})
    chairman_initials = chairman.get('name_initials', '—')
    vice_chairman = commission_data.get('vice_chairman', {})
    vice_chairman_initials = vice_chairman.get('name_initials', '—')

    if commission:
        if commission.department:
            binding = decline_phrase(commission.department.name, 'gent')
        elif commission.subdivision:
            binding = decline_phrase(commission.subdivision.name, 'gent')
        elif commission.organization:
            binding = commission.organization.short_name_ru
        else:
            binding = ""
    else:
        binding = ""

    context.setdefault('chairman_name_initials', chairman_initials)
    context.setdefault('vice_chairman_name_initials', vice_chairman_initials)
    context.setdefault('binding_name_genitive', binding)

    if commission_cache is not None:
        commission_cache[cache_key] = {
            'chairman_name_initials': chairman_initials,
            'vice_chairman_name_initials': vice_chairman_initials,
            'binding_name_genitive': binding,
        }
    return context


def _trim_rows(table, keep_count):
    for idx in range(len(table.rows) - 1, keep_count - 1, -1):
        table._tbl.remove(table.rows[idx]._tr)


def _empty_employee_slot():
    return {
        'fio_dative': '',
        'position_nominative': '',
        'chairman_name_initials': '',
        'vice_chairman_name_initials': '',
        'binding_or_org_genitive': '',
        'organization_name': '',
    }


def _build_employee_slot(employee, commission_cache: Optional[Dict[int, Dict[str, Any]]] = None) -> Dict[str, Any]:
    context = _build_certificate_context(employee, commission_cache=commission_cache)
    binding_or_org = context.get('binding_name_genitive') or context.get('organization_name_genitive') or ''
    workplace = ''
    if employee.department:
        workplace = employee.department.name
    elif employee.subdivision:
        workplace = employee.subdivision.name
    return {
        'fio_dative': context.get('fio_dative', ''),
        'position_nominative': context.get('position_nominative', ''),
        'chairman_name_initials': context.get('chairman_name_initials', ''),
        'vice_chairman_name_initials': context.get('vice_chairman_name_initials', ''),
        'binding_or_org_genitive': binding_or_org,
        'organization_name': context.get('organization_name', ''),
        'workplace': workplace,
    }


def generate_safety_certificates(
    employees: List[Any],
    grouping_name: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    if not employees:
        return None

    if not TEMPLATE_PATH.exists():
        logger.error("Шаблон удостоверений не найден: %s", TEMPLATE_PATH)
        return None

    template_bytes = _get_template_bytes()
    groups = [employees[i:i + 4] for i in range(0, len(employees), 4)]
    composer = None

    try:
        commission_cache = {}
        for group in groups:
            render_contexts = [_build_employee_slot(employee, commission_cache=commission_cache) for employee in group]
            while len(render_contexts) < 4:
                render_contexts.append(_empty_employee_slot())

            template = DocxTemplate(BytesIO(template_bytes))
            template.render({'employees': render_contexts})
            doc = template.docx

            if len(doc.tables) < 2:
                logger.error("Шаблон удостоверений должен содержать две таблицы (лицевая/оборотная)")
                return None

            front_table = doc.tables[0]
            back_table = doc.tables[1]

            if len(group) < 4:
                _trim_rows(front_table, len(group))
                _trim_rows(back_table, len(group))

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            rendered_doc = Document(buffer)

            if composer is None:
                composer = Composer(rendered_doc)
            else:
                composer.append(rendered_doc)
    except Exception:
        logger.error("Ошибка формирования удостоверений", exc_info=True)
        return None

    if composer is None:
        return None

    buffer = BytesIO()
    composer.save(buffer)
    buffer.seek(0)

    if grouping_name:
        clean_name = grouping_name.replace('"', '').replace("'", '').replace('«', '').replace('»', '')
        filename = f"Удостоверения по ОТ_{clean_name}.docx"
    else:
        org_name = employees[0].organization.short_name_ru if employees[0].organization else "Организация"
        clean_name = org_name.replace('"', '').replace("'", '').replace('«', '').replace('»', '')
        filename = f"Удостоверения по ОТ_{clean_name}.docx"

    return {'content': buffer.getvalue(), 'filename': filename}
