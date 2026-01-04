"""
Генератор удостоверений по охране труда с построчным рендерингом.

Оригинальная концепция:
- Для каждого сотрудника рендерится ОДНА строка отдельно через docxtpl
- Отрендеренные строки копируются в финальный документ
- Группы по 5 сотрудников на страницу (лицевая + оборотная стороны)
- Склеивание через docxcompose.Composer без прямого XML манипулирования
"""
import logging
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

from docx import Document
from docxcompose.composer import Composer
from docxtpl import DocxTemplate

from directory.document_generators.base import prepare_employee_context
from directory.utils import find_appropriate_commission, get_commission_members_formatted
from directory.utils.declension import decline_phrase
from directory.utils.docx_cleaner import clean_document

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(
    "/home/django/webapps/potby/media/document_templates/etalon/udostoverneie_blank.docx"
)


def _trim_rows(table, keep_count):
    """
    Удаляет лишние строки из таблицы.

    Использует безопасный метод удаления через _tbl.remove() без повреждения XML.

    Args:
        table: Таблица docx
        keep_count: Количество строк, которые нужно оставить
    """
    for idx in range(len(table.rows) - 1, keep_count - 1, -1):
        table._tbl.remove(table.rows[idx]._tr)


def _build_employee_context(employee, commission_cache: Optional[Dict[int, Dict[str, Any]]] = None) -> Dict[str, Any]:
    """
    Строит контекст для рендеринга удостоверения одного сотрудника.

    Использует существующие функции:
    - prepare_employee_context() из base.py
    - find_appropriate_commission() из utils
    - get_commission_members_formatted() из utils

    Args:
        employee: Объект сотрудника
        commission_cache: Кэш для оптимизации запросов к комиссиям

    Returns:
        Словарь с контекстом для шаблона:
        - fio_dative: ФИО в дательном падеже
        - position_nominative: Должность
        - chairman_name_initials: Инициалы председателя комиссии
        - vice_chairman_name_initials: Инициалы заместителя председателя комиссии
        - binding_name_genitive: Название подразделения/отдела в родительном падеже
        - organization_name_genitive: Название организации в родительном падеже
        - organization_name: Название организации
        - workplace: Место работы (2 нижних уровня иерархии)
    """
    context = prepare_employee_context(employee)
    commission = find_appropriate_commission(employee)

    # Кэширование для оптимизации
    cache_key = commission.id if commission else None
    if commission_cache is not None and cache_key in commission_cache:
        cached = commission_cache[cache_key]
        # Объединяем с базовым контекстом сотрудника
        result = context.copy()
        result.update(cached)

        # Вычисляем "Место работы" (2 нижних уровня иерархии)
        # Это поле зависит от сотрудника, поэтому не кэшируется
        if context.get('department'):
            workplace = context.get('department', '')
        elif context.get('subdivision'):
            workplace = context.get('subdivision', '')
        else:
            workplace = context.get('organization_name', '')
        result['workplace'] = workplace

        return result

    # Получаем данные комиссии
    commission_data = get_commission_members_formatted(commission) if commission else {}
    chairman = commission_data.get('chairman', {})
    chairman_initials = chairman.get('name_initials', '')
    vice_chairman = commission_data.get('vice_chairman', {})
    vice_chairman_initials = vice_chairman.get('name_initials', '')

    # Определяем binding (привязку к подразделению/отделу)
    if commission:
        if commission.department:
            binding = decline_phrase(commission.department.name, 'gent')
        elif commission.subdivision:
            binding = decline_phrase(commission.subdivision.name, 'gent')
        elif commission.organization:
            binding = commission.organization.short_name_ru
        else:
            binding = ""
    else:
        binding = ""

    # Формируем переменные для шаблона
    # Шаблон использует binding_name_genitive и organization_name_genitive

    # Вычисляем "Место работы" (2 нижних уровня иерархии)
    if context.get('department'):
        workplace = context.get('department', '')
    elif context.get('subdivision'):
        workplace = context.get('subdivision', '')
    else:
        workplace = context.get('organization_name', '')

    result = {
        'fio_dative': context.get('fio_dative', ''),
        'position_nominative': context.get('position_nominative', ''),
        'chairman_name_initials': chairman_initials,
        'vice_chairman_name_initials': vice_chairman_initials,
        'binding_name_genitive': binding if binding else '',
        'organization_name_genitive': context.get('organization_name_genitive', ''),
        'organization_name': context.get('organization_name', ''),
        'workplace': workplace,
    }

    # Сохраняем в кэш
    if commission_cache is not None:
        commission_cache[cache_key] = {
            'chairman_name_initials': chairman_initials,
            'vice_chairman_name_initials': vice_chairman_initials,
            'binding_name_genitive': result['binding_name_genitive'],
            'organization_name_genitive': result['organization_name_genitive'],
            'organization_name': result['organization_name'],
        }

    return result


def _render_employee_rows(template_bytes: bytes, employee_context: Dict) -> tuple:
    """
    Рендерит строки удостоверения для одного сотрудника.

    Алгоритм:
    1. Загружаем шаблон через DocxTemplate
    2. Рендерим ВЕСЬ документ с контекстом сотрудника
    3. Извлекаем первую строку из каждой таблицы (она отрендерена)
    4. Возвращаем эти строки

    Args:
        template_bytes: Байты шаблона docx
        employee_context: Контекст сотрудника для рендеринга

    Returns:
        tuple: (front_row, back_row) - строки для лицевой и оборотной стороны
    """
    # Создаем временный документ
    template = DocxTemplate(BytesIO(template_bytes))

    # Рендерим с контекстом
    template.render(employee_context)

    # Получаем отрендеренный документ
    doc = template.docx

    # Извлекаем первую строку из каждой таблицы
    front_row = doc.tables[0].rows[0]
    back_row = doc.tables[1].rows[0]

    return front_row, back_row


def _copy_cell_content(source_cell, target_cell):
    """
    Копирует содержимое и форматирование ячейки.

    Безопасный метод без прямого XML манипулирования.
    Копирует:
    - Текст
    - Форматирование параграфов (выравнивание, межстрочный интервал)
    - Форматирование runs (шрифт, размер, жирность, курсив, подчёркивание)

    Args:
        source_cell: Исходная ячейка
        target_cell: Целевая ячейка
    """
    # Очищаем целевую ячейку
    for paragraph in target_cell.paragraphs:
        paragraph.clear()

    # Удаляем все параграфы кроме первого
    while len(target_cell.paragraphs) > 1:
        target_cell._element.remove(target_cell.paragraphs[-1]._element)

    # Копируем параграфы из источника
    first = True
    for paragraph in source_cell.paragraphs:
        # Используем существующий первый параграф
        if first:
            target_para = target_cell.paragraphs[0]
            first = False
        else:
            target_para = target_cell.add_paragraph()

        # Копируем стиль параграфа, чтобы сохранить размер шрифта по стилю
        if paragraph.style is not None:
            target_para.style = paragraph.style

        # Копируем форматирование параграфа
        target_para.alignment = paragraph.alignment

        # Копируем межстрочный интервал
        if paragraph.paragraph_format.line_spacing is not None:
            target_para.paragraph_format.line_spacing = paragraph.paragraph_format.line_spacing

        # Копируем отступы (даже если они равны 0)
        if paragraph.paragraph_format.space_before is not None:
            target_para.paragraph_format.space_before = paragraph.paragraph_format.space_before
        if paragraph.paragraph_format.space_after is not None:
            target_para.paragraph_format.space_after = paragraph.paragraph_format.space_after

        # Копируем отступ первой строки
        if paragraph.paragraph_format.first_line_indent is not None:
            target_para.paragraph_format.first_line_indent = paragraph.paragraph_format.first_line_indent

        # Копируем отступы слева/справа
        if paragraph.paragraph_format.left_indent is not None:
            target_para.paragraph_format.left_indent = paragraph.paragraph_format.left_indent
        if paragraph.paragraph_format.right_indent is not None:
            target_para.paragraph_format.right_indent = paragraph.paragraph_format.right_indent

        # Копируем runs
        for run in paragraph.runs:
            new_run = target_para.add_run(run.text)
            if run.style is not None:
                new_run.style = run.style
            if run.font.name:
                new_run.font.name = run.font.name
            if run.font.size:
                new_run.font.size = run.font.size
            if run.font.bold is not None:
                new_run.font.bold = run.font.bold
            if run.font.italic is not None:
                new_run.font.italic = run.font.italic
            if run.font.underline is not None:
                new_run.font.underline = run.font.underline
            if run.font.color is not None and run.font.color.rgb is not None:
                new_run.font.color.rgb = run.font.color.rgb
            if run.font.highlight_color is not None:
                new_run.font.highlight_color = run.font.highlight_color
            if run.font.all_caps is not None:
                new_run.font.all_caps = run.font.all_caps
            if run.font.small_caps is not None:
                new_run.font.small_caps = run.font.small_caps
            if run.font.strike is not None:
                new_run.font.strike = run.font.strike
            if run.font.double_strike is not None:
                new_run.font.double_strike = run.font.double_strike
            if run.font.subscript is not None:
                new_run.font.subscript = run.font.subscript
            if run.font.superscript is not None:
                new_run.font.superscript = run.font.superscript
            if run.font.shadow is not None:
                new_run.font.shadow = run.font.shadow
            if run.font.outline is not None:
                new_run.font.outline = run.font.outline


def _assemble_document(template_bytes: bytes, employees: List, commission_cache: Dict) -> Document:
    """
    Собирает финальный документ для группы сотрудников (до 5 человек).

    Алгоритм:
    1. Рендерим строки для каждого сотрудника
    2. Загружаем чистый шаблон
    3. Копируем отрендеренные строки в соответствующие позиции
    4. Удаляем лишние строки если сотрудников < 5

    Args:
        template_bytes: Байты шаблона docx
        employees: Список сотрудников (до 5 человек)
        commission_cache: Кэш комиссий для оптимизации

    Returns:
        Document: Готовый документ с заполненными удостоверениями
    """
    # Рендерим строки для всех сотрудников
    rendered_rows = []
    for employee in employees:
        context = _build_employee_context(employee, commission_cache)
        front_row, back_row = _render_employee_rows(template_bytes, context)
        rendered_rows.append((front_row, back_row))

    # Загружаем чистый шаблон для финального документа
    final_doc = Document(BytesIO(template_bytes))

    # Копируем отрендеренные строки
    for idx, (front_row, back_row) in enumerate(rendered_rows):
        if idx < len(final_doc.tables[0].rows):
            # Копируем в лицевую таблицу
            target_front_row = final_doc.tables[0].rows[idx]
            for cell_idx in range(min(len(front_row.cells), len(target_front_row.cells))):
                _copy_cell_content(front_row.cells[cell_idx], target_front_row.cells[cell_idx])

            # Копируем в оборотную таблицу
            target_back_row = final_doc.tables[1].rows[idx]
            for cell_idx in range(min(len(back_row.cells), len(target_back_row.cells))):
                _copy_cell_content(back_row.cells[cell_idx], target_back_row.cells[cell_idx])

    # Удаляем лишние строки
    if len(employees) < 5:
        _trim_rows(final_doc.tables[0], len(employees))
        _trim_rows(final_doc.tables[1], len(employees))

    return final_doc


def generate_safety_certificates_rowwise(
    employees: List[Any],
    grouping_name: Optional[str] = None
) -> Optional[Dict[str, Any]]:
    """
    Генерирует удостоверения по охране труда через построчный рендеринг.

    Оригинальная концепция:
    - Каждая строка рендерится отдельно через docxtpl
    - Группы по 5 сотрудников на страницу
    - Склеивание через Composer без прямого XML

    Args:
        employees: Список сотрудников для генерации удостоверений
        grouping_name: Название группы для имени файла (необязательно)

    Returns:
        Dict с ключами:
        - content: байты сгенерированного docx файла
        - filename: название файла
        Или None в случае ошибки
    """
    # Проверки безопасности
    if not employees:
        logger.warning("Пустой список сотрудников")
        return None

    if not TEMPLATE_PATH.exists():
        logger.error(f"Шаблон не найден: {TEMPLATE_PATH}")
        return None

    template_bytes = TEMPLATE_PATH.read_bytes()

    # Проверка структуры шаблона
    test_doc = Document(BytesIO(template_bytes))
    if len(test_doc.tables) != 2:
        logger.error("Шаблон должен содержать 2 таблицы (лицевая и оборотная стороны)")
        return None
    if len(test_doc.tables[0].rows) != 5 or len(test_doc.tables[1].rows) != 5:
        logger.error("Каждая таблица должна содержать 5 строк")
        return None

    # Разбиваем на группы по 5 сотрудников
    groups = [employees[i:i + 5] for i in range(0, len(employees), 5)]
    composer = None
    commission_cache = {}

    try:
        for group in groups:
            # Рендерим документ для группы
            doc = _assemble_document(template_bytes, group, commission_cache)

            # Склеиваем через Composer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            rendered_doc = Document(buffer)

            if composer is None:
                composer = Composer(rendered_doc)
            else:
                composer.append(rendered_doc)

    except Exception:
        logger.error("Ошибка построчного рендеринга удостоверений", exc_info=True)
        return None

    # Сохраняем результат
    if composer is None:
        return None

    buffer = BytesIO()
    composer.save(buffer)
    buffer.seek(0)

    # Очищаем пустые параграфы и строки
    cleaned_content = clean_document(buffer.getvalue())

    # Формируем имя файла
    if grouping_name:
        clean_name = grouping_name.replace('"', '').replace("'", '').replace('«', '').replace('»', '')
        filename = f"Удостоверения по ОТ_{clean_name}.docx"
    else:
        org_name = employees[0].organization.short_name_ru if employees[0].organization else "Организация"
        clean_name = org_name.replace('"', '').replace("'", '').replace('«', '').replace('»', '')
        filename = f"Удостоверения по ОТ_{clean_name}.docx"

    return {'content': cleaned_content, 'filename': filename}
