# directory/document_generators/equipment_journal_generator.py
"""
📄 Генератор для журнала периодического осмотра оборудования (грузовые тележки и др.)
"""
import logging
import traceback
import datetime
from typing import Dict, Any, Optional, List
from io import BytesIO
from docxtpl import DocxTemplate
from docx.shared import Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from directory.document_generators.base import get_document_template
from directory.utils.declension import get_initials_from_name

# Настройка логирования
logger = logging.getLogger(__name__)

LABEL_EQUIPMENT_TYPES = {
    'Грузовая тележка',
    'Грузовые тележки',
}


_LADDER_TYPE_ID = None


def _get_ladder_type_id():
    global _LADDER_TYPE_ID
    if _LADDER_TYPE_ID is not None:
        return _LADDER_TYPE_ID
    try:
        from deadline_control.models import EquipmentType
        _LADDER_TYPE_ID = EquipmentType.objects.filter(name='Лестница').values_list('id', flat=True).first()
    except Exception:
        _LADDER_TYPE_ID = None
    return _LADDER_TYPE_ID


def _resolve_template_code(equipment_type=None, equipment_type_name: str = '') -> str:
    ladder_type_id = _get_ladder_type_id()
    if equipment_type and ladder_type_id and equipment_type.id == ladder_type_id:
        return 'lestnicy-journal'
    if equipment_type_name:
        normalized = equipment_type_name.strip().lower()
        if normalized == 'лестница':
            return 'lestnicy-journal'
    return 'equipment_journal'


def _sanitize_filename(value: str) -> str:
    """
    Удаляет недопустимые символы для имени файла.
    """
    if not value:
        return ''
    invalid_chars = ['"', '/', '\\', ':', '*', '?', '<', '>', '|']
    for ch in invalid_chars:
        value = value.replace(ch, '_')
    return value.strip()


def _get_equipment_label(equipment_type_name: str) -> str:
    """
    Возвращает корректное название типа оборудования для имени файла.
    """
    if not equipment_type_name:
        return ''
    normalized = equipment_type_name.strip()
    lower_name = normalized.lower()
    if lower_name in ['грузовая тележка', 'грузовые тележки']:
        return 'тележек'
    return normalized


def _find_equipment_journal_table(docx_doc):
    """
    Находит таблицу журнала оборудования в документе.
    Обычно это последняя (или единственная) таблица в документе.
    """
    return docx_doc.tables[-1] if docx_doc.tables else None


def _set_cell_borders(cell):
    """
    Устанавливает границы для ячейки таблицы.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.oxml.ns import qn

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Удаляем старые границы, если есть
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)

    # Создаём элемент границ
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tcBorders>
    ''')

    # Добавляем новые границы
    tcPr.append(tcBorders)


def _reset_equipment_journal_table(table):
    """
    Удаляет все строки данных, оставляя только заголовки.
    Находит последнюю строку заголовка по первой ячейке (должна содержать "1").
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    # Ищем последнюю строку заголовка - проверяем первую ячейку каждой строки
    last_header_row_idx = 1  # По умолчанию предполагаем, что это вторая строка (индекс 1)

    for row_idx, row in enumerate(table.rows):
        if len(row.cells) > 0:
            first_cell_text = row.cells[0].text.strip()
            logger.info(f"Проверка строки {row_idx}: первая ячейка = '{first_cell_text}'")
            # Если первая ячейка содержит "1" - это строка с номерами столбцов
            if (first_cell_text == '1' or
                first_cell_text == '1.' or
                first_cell_text == '1)' or
                first_cell_text.startswith('1')):
                last_header_row_idx = row_idx
                logger.info(f"✓ Найдена строка с номерами столбцов: индекс {row_idx}")
                break

    # Количество строк заголовка = индекс строки с номерами + 1
    num_header_rows = last_header_row_idx + 1
    logger.info(f"Количество строк заголовка: {num_header_rows}")

    # Удаляем ВСЕ строки после заголовков
    while len(table.rows) > num_header_rows:
        row = table.rows[-1]
        tbl = table._tbl
        tbl.remove(row._tr)

    # Устанавливаем границы для всех ячеек в строках заголовка
    for row_idx in range(num_header_rows):
        for cell in table.rows[row_idx].cells:
            _set_cell_borders(cell)

    # Устанавливаем повторение заголовков на каждой странице
    from docx.oxml.ns import qn

    for row_idx in range(num_header_rows):
        header_row = table.rows[row_idx]
        tr = header_row._tr
        trPr = tr.trPr

        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
            tr.insert(0, trPr)

        # Проверяем, нет ли уже тега tblHeader
        existing_tblHeader = trPr.find(qn('w:tblHeader'))
        if existing_tblHeader is not None:
            trPr.remove(existing_tblHeader)

        # Добавляем тег tblHeader для повторения этой строки
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trPr.insert(0, tblHeader)
        logger.info(f"  ✓ Установлено повторение для строки {row_idx}")


def _fill_equipment_journal_rows(table, equipment_records: List[Dict[str, str]], rows_per_page: int = 24):
    """
    Заполняет таблицу журнала оборудования строками с данными и пустыми строками до конца страницы.

    Args:
        table: Таблица документа Word
        equipment_records: Данные оборудования для заполнения
            Каждая запись должна содержать:
            - type: Тип оборудования
            - inventory_number: Инвентарный номер
            - location: Место эксплуатации
            - inspection_date: Дата осмотра (строка в формате ДД.ММ.ГГГГ)
            - inspector: Должность, ФИО проверяющего (может быть пустым)
            - result: Результат осмотра (может быть пустым)
        rows_per_page: Количество строк на одной странице (по умолчанию 30)
    """
    logger.info(f"Добавление {len(equipment_records)} записей оборудования в таблицу")

    # Находим последнюю строку заголовка (строку с номерами столбцов)
    header_row_idx = 1  # По умолчанию
    for row_idx, row in enumerate(table.rows):
        if len(row.cells) > 0:
            first_cell_text = row.cells[0].text.strip()
            if (first_cell_text == '1' or
                first_cell_text == '1.' or
                first_cell_text.startswith('1')):
                header_row_idx = row_idx
                break

    # Получаем количество столбцов из строки с номерами
    num_cols = len(table.rows[header_row_idx].cells)
    logger.info(f"Количество столбцов в таблице: {num_cols}")

    # Добавляем строки с данными
    for idx, record in enumerate(equipment_records, start=1):
        # Добавляем новую строку
        new_row = table.add_row()
        logger.info(f"Добавлена строка {idx}/{len(equipment_records)}")

        # Заполняем ячейки
        cells = new_row.cells

        # Столбец 1: № п/п (сквозная нумерация)
        cells[0].text = str(idx)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 2: Тип оборудования
        cells[1].text = record.get('type', '')
        cells[1].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 3: Инвентарный номер
        inventory_number = record.get('inventory_number', '')
        if inventory_number:
            cells[2].text = f"Инв. № {inventory_number}"
        else:
            cells[2].text = ''
        cells[2].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 4: Место эксплуатации
        cells[3].text = record.get('location', '')
        cells[3].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 5: Дата осмотра
        cells[4].text = record.get('inspection_date', '')
        cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[4].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 6: Должность, ФИО проверяющего (пусто для ручного заполнения)
        cells[5].text = record.get('inspector', '')
        cells[5].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 7: Результаты осмотра (пусто для ручного заполнения)
        cells[6].text = record.get('result', '')
        cells[6].paragraphs[0].runs[0].font.size = Pt(14)

        # Столбец 8: Подпись (всегда пустое)
        cells[7].text = ''

        # Устанавливаем границы для всех ячеек строки
        for cell in cells:
            _set_cell_borders(cell)

    logger.info(f"✓ Успешно добавлено {len(equipment_records)} записей")

    # Пустые строки не добавляем


def _fill_ladder_journal_rows(table, ladder_records: List[Dict[str, str]]):
    """
    Заполняет таблицу журнала испытаний лестниц.

    Ожидаемые поля записи:
    - object_name
    - inspection_date
    - inventory_number
    - result
    - next_inspection_date
    """
    logger.info(f"Добавление {len(ladder_records)} записей лестниц в таблицу")

    header_row_idx = 1
    for row_idx, row in enumerate(table.rows):
        if len(row.cells) > 0:
            first_cell_text = row.cells[0].text.strip()
            if (first_cell_text == '1' or
                first_cell_text == '1.' or
                first_cell_text.startswith('1')):
                header_row_idx = row_idx
                break

    num_cols = len(table.rows[header_row_idx].cells)
    if num_cols < 7:
        logger.error(f"Недостаточно столбцов для журнала лестниц: {num_cols}")
        return

    for idx, record in enumerate(ladder_records, start=1):
        new_row = table.add_row()
        cells = new_row.cells

        cells[0].text = str(idx)
        cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[0].paragraphs[0].runs[0].font.size = Pt(14)

        cells[1].text = record.get('object_name', '')
        cells[1].paragraphs[0].runs[0].font.size = Pt(14)

        cells[2].text = record.get('inspection_date', '')
        cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[2].paragraphs[0].runs[0].font.size = Pt(14)

        inventory_number = record.get('inventory_number', '')
        cells[3].text = inventory_number or ''
        cells[3].paragraphs[0].runs[0].font.size = Pt(14)

        cells[4].text = record.get('result', '')
        cells[4].paragraphs[0].runs[0].font.size = Pt(14)

        cells[5].text = record.get('next_inspection_date', '')
        cells[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cells[5].paragraphs[0].runs[0].font.size = Pt(14)

        cells[6].text = ''

        for cell in cells:
            _set_cell_borders(cell)

    logger.info(f"? Успешно добавлено {len(ladder_records)} записей")

def _post_process_equipment_journal(doc, context: Dict[str, Any]):
    """
    Пост-обработчик документа: очищает и заполняет таблицу журнала.

    Args:
        doc: DocxTemplate документ
        context: Контекст с данными для заполнения

    Returns:
        DocxTemplate: Обработанный документ
    """
    logger.info("Начало пост-обработки документа журнала оборудования")

    # Находим таблицу
    table = _find_equipment_journal_table(doc)
    if not table:
        logger.error("Таблица не найдена в документе")
        return doc

    logger.info(f"Найдена таблица с {len(table.rows)} строками")

    # Очищаем таблицу от строк данных
    _reset_equipment_journal_table(table)

    # Получаем данные оборудования
    template_code = context.get('template_code')
    if template_code == 'lestnicy-journal':
        ladder_records = context.get('ladder_records', [])
        if not ladder_records:
            logger.warning("Нет записей лестниц для добавления в журнал")
            return doc
        _fill_ladder_journal_rows(table, ladder_records)
    else:
        equipment_records = context.get('equipment_records', [])
        if not equipment_records:
            logger.warning("Нет записей оборудования для добавления в журнал")
            return doc
        _fill_equipment_journal_rows(table, equipment_records)

    logger.info("✓ Пост-обработка документа завершена")
    return doc


def _append_equipment_labels(doc, equipment_list):
    """
    Добавляет лист с бирками для оборудования.
    """
    if not equipment_list:
        return

    doc.add_page_break()
    heading = doc.add_paragraph("Бирки для тележек")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(20)

    description = doc.add_paragraph(
        "Бирки нужно вырезать и наклеить на соответствующие тележки."
    )
    description.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in description.runs:
        run.font.size = Pt(15)

    table = doc.add_table(rows=0, cols=2)

    for idx, equipment in enumerate(equipment_list):
        if idx % 2 == 0:
            row = table.add_row()

        cell = row.cells[idx % 2]
        cell.text = ''
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        capacity = equipment.load_capacity_kg
        location_parts = []
        if equipment.organization:
            location_parts.append(equipment.organization.short_name_ru)
        if equipment.subdivision:
            location_parts.append(equipment.subdivision.name)
        if equipment.department:
            location_parts.append(equipment.department.name)

        if len(location_parts) >= 2:
            subdivision_label = " / ".join(location_parts[-2:])
        elif location_parts:
            subdivision_label = location_parts[0]
        else:
            subdivision_label = 'Без подразделения'

        lines = []
        lines.append(f"{subdivision_label}")
        lines.append(f"{equipment.equipment_name}")
        if capacity is not None:
            lines.append(f"Грузоподъемность: {capacity} кг")
        lines.append(f"Инв. № {equipment.inventory_number}")

        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line_idx, line in enumerate(lines):
            if line_idx > 0:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.size = Pt(14)

        _set_cell_borders(cell)


def _append_ladder_labels(doc, equipment_list, inspection_date=None):
    """
    Добавляет лист с бирками для лестниц.
    """
    if not equipment_list:
        return

    doc.add_page_break()
    heading = doc.add_paragraph("Бирки для лестниц")
    heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.size = Pt(20)

    table = doc.add_table(rows=0, cols=2)

    for idx, equipment in enumerate(equipment_list):
        if idx % 2 == 0:
            row = table.add_row()

        cell = row.cells[idx % 2]
        cell.text = ''
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        organization_name = equipment.organization.short_name_ru if equipment.organization else ''
        subdivision_name = equipment.subdivision.name if equipment.subdivision else ''
        location_parts = [part for part in [organization_name, subdivision_name] if part]
        location_label = " / ".join(location_parts) if location_parts else "Без подразделения"

        next_date = ''
        resolved_date = _resolve_inspection_date(inspection_date)
        if resolved_date and equipment.maintenance_period_months:
            next_date = _add_months(resolved_date, equipment.maintenance_period_months).strftime('%d.%m.%Y')

        lines = [
            location_label,
            f"Тип (марка): {equipment.equipment_name}",
            f"Инв. № {equipment.inventory_number}",
            f"Следующее испытание: {next_date}" if next_date else "Следующее испытание: -",
        ]

        paragraph = cell.paragraphs[0]
        paragraph.clear()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for line_idx, line in enumerate(lines):
            if line_idx > 0:
                paragraph.add_run().add_break()
            run = paragraph.add_run(line)
            run.font.size = Pt(12)

        _set_cell_borders(cell)


def _resolve_inspection_date(inspection_date):
    if isinstance(inspection_date, datetime.date):
        return inspection_date
    if isinstance(inspection_date, str):
        try:
            return datetime.date.fromisoformat(inspection_date)
        except ValueError:
            return None
    return None


def _add_months(base_date, months):
    if not base_date or months is None:
        return base_date
    year = base_date.year + (base_date.month - 1 + months) // 12
    month = (base_date.month - 1 + months) % 12 + 1
    import calendar
    last_day = calendar.monthrange(year, month)[1]
    day = min(base_date.day, last_day)
    return datetime.date(year, month, day)


def _build_equipment_records(equipment_list, inspection_date=None, use_two_level_location=False):
    records = []
    resolved_date = _resolve_inspection_date(inspection_date)
    for eq in equipment_list:
        if use_two_level_location:
            location_parts = []
            if eq.organization:
                location_parts.append(eq.organization.short_name_ru)
            if eq.subdivision:
                location_parts.append(eq.subdivision.name)
            if eq.department:
                location_parts.append(eq.department.name)

            if len(location_parts) >= 2:
                location = " / ".join(location_parts[-2:])
            elif location_parts:
                location = location_parts[0]
            else:
                location = ''
        else:
            if eq.department:
                location = eq.department.name
            elif eq.subdivision:
                location = eq.subdivision.name
            elif eq.organization:
                location = eq.organization.short_name_ru
            else:
                location = ''

        eq_type = eq.equipment_type.name if eq.equipment_type else eq.equipment_name

        if resolved_date:
            inspection_date_str = resolved_date.strftime('%d.%m.%Y')
        elif eq.last_maintenance_date:
            inspection_date_str = eq.last_maintenance_date.strftime('%d.%m.%Y')
        else:
            inspection_date_str = ''

        records.append({
            'type': eq_type,
            'inventory_number': eq.inventory_number,
            'location': location,
            'inspection_date': inspection_date_str,
            'inspector': '',
            'result': '',
        })

    return records


def _build_ladder_records(equipment_list, inspection_date=None):
    records = []
    resolved_date = _resolve_inspection_date(inspection_date)
    for eq in equipment_list:
        if resolved_date:
            inspection_date_str = resolved_date.strftime('%d.%m.%Y')
        elif eq.last_maintenance_date:
            inspection_date_str = eq.last_maintenance_date.strftime('%d.%m.%Y')
        else:
            inspection_date_str = ''

        if eq.next_maintenance_date:
            next_date_str = eq.next_maintenance_date.strftime('%d.%m.%Y')
        else:
            next_date_str = ''

        records.append({
            'object_name': eq.equipment_name,
            'inspection_date': inspection_date_str,
            'inventory_number': eq.inventory_number,
            'result': '',
            'next_inspection_date': next_date_str,
        })

    return records


def generate_equipment_journal_for_subdivision(
    equipment,
    equipment_type,
    inspection_date,
    subdivision,
    subdivision_name=None,
    use_two_level_location=False
):
    """
    Генерирует журнал осмотра оборудования для подразделения.
    """
    try:
        template_code = _resolve_template_code(equipment_type=equipment_type)
        template = get_document_template(template_code, employee=None)
        if not template:
            logger.error(f"Шаблон документа '{template_code}' не найден")
            return None

        equipment_list = list(equipment)
        if not equipment_list:
            logger.warning("Нет оборудования для генерации журнала")
            return None

        resolved_date = _resolve_inspection_date(inspection_date) or datetime.date.today()
        year_start = datetime.date(resolved_date.year, 1, 1)
        year_end = datetime.date(resolved_date.year, 12, 31)

        organization = None
        if subdivision:
            organization = subdivision.organization
        elif equipment_list and equipment_list[0].organization:
            organization = equipment_list[0].organization

        if organization:
            org_full = organization.full_name_ru
            org_short = organization.short_name_ru
        else:
            org_full = "Все организации"
            org_short = "Все организации"

        if subdivision:
            structural_unit = subdivision.name
        elif organization:
            structural_unit = organization.full_name_ru
        else:
            structural_unit = org_full

        if template_code == 'lestnicy-journal':
            ladder_records = _build_ladder_records(
                equipment_list,
                inspection_date=inspection_date
            )
            equipment_records = []
        else:
            ladder_records = []
            equipment_records = _build_equipment_records(
                equipment_list,
                inspection_date=inspection_date,
                use_two_level_location=use_two_level_location
            )

        context = {
            'organization': {
                'full_name_ru': org_full,
                'short_name_ru': org_short,
            },
            'start_date': year_start,
            'end_date': year_end,
            'equipment_records': equipment_records,
            'ladder_records': ladder_records,
            'template_code': template_code,
            'subdivision_name': subdivision.name if subdivision else (subdivision_name or ''),
            'structural_unit': structural_unit,
        }

        template_path = template.template_file.path
        doc = DocxTemplate(template_path)
        doc.render(context)
        doc = _post_process_equipment_journal(doc, context)

        if equipment_type.name in LABEL_EQUIPMENT_TYPES:
            _append_equipment_labels(doc, equipment_list)
        if template_code == 'lestnicy-journal':
            _append_ladder_labels(doc, equipment_list, inspection_date=resolved_date)

        name_part = subdivision.name if subdivision else (subdivision_name or "Общий")
        safe_name = _sanitize_filename(name_part)
        label = _get_equipment_label(equipment_type.name)
        date_str = resolved_date.strftime('%d.%m.%Y')
        if safe_name:
            filename = f"Журнал осмотра {label} {safe_name} {date_str}.docx"
        else:
            filename = f"Журнал осмотра {label} {date_str}.docx"

        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        return {
            'content': docx_buffer.getvalue(),
            'filename': filename,
        }
    except Exception as exc:
        logger.error(f"Ошибка при генерации журнала оборудования: {str(exc)}")
        logger.error(traceback.format_exc())
        return None


def generate_equipment_journal(organization, equipment_type_name, start_date=None, end_date=None, inspection_date=None):
    """
    Генерирует журнал периодического осмотра оборудования.

    Args:
        organization: Объект Organization
        equipment_type_name: Название типа оборудования (например, "Грузовая тележка")
        start_date: Дата начала журнала (по умолчанию - начало текущего года)
        end_date: Дата окончания журнала (по умолчанию - конец текущего года)

    Returns:
        Dict с 'content' (байты файла) и 'filename' или None при ошибке
    """
    try:
        from deadline_control.models import Equipment

        logger.info(f"Генерация журнала для организации: {organization.short_name_ru}")
        logger.info(f"Тип оборудования: {equipment_type_name}")

        # Устанавливаем даты по умолчанию
        if not start_date:
            start_date = datetime.date(datetime.datetime.now().year, 1, 1)
        if not end_date:
            end_date = datetime.date(datetime.datetime.now().year, 12, 31)

        logger.info(f"Период журнала: {start_date} - {end_date}")

        # Получаем шаблон документа
        template_code = _resolve_template_code(equipment_type_name=equipment_type_name)
        template = get_document_template(template_code, employee=None)
        if not template:
            logger.error(f"Шаблон документа '{template_code}' не найден")
            return None

        # Получаем оборудование из базы данных
        equipment_list = Equipment.objects.filter(
            equipment_type__name__icontains=equipment_type_name,
            organization=organization
        ).select_related(
            'equipment_type', 'organization', 'subdivision', 'department'
        ).order_by('inventory_number')

        logger.info(f"Найдено оборудования: {equipment_list.count()}")

        if template_code == 'lestnicy-journal':
            ladder_records = _build_ladder_records(
                equipment_list,
                inspection_date=inspection_date
            )
            equipment_records = []
        else:
            ladder_records = []
            equipment_records = _build_equipment_records(
                equipment_list,
                inspection_date=inspection_date,
                use_two_level_location=False
            )

        logger.info(f"Подготовлено {len(equipment_records)} записей для журнала")

        # Формируем контекст для шаблона
        context = {
            'organization': {
                'full_name_ru': organization.full_name_ru,
                'short_name_ru': organization.short_name_ru,
            },
            'start_date': start_date,
            'end_date': end_date,
            'equipment_records': equipment_records,
            'ladder_records': ladder_records,
            'template_code': template_code,
            'structural_unit': organization.full_name_ru,
        }

        # Генерируем документ
        template_path = template.template_file.path
        logger.info(f"Загрузка шаблона: {template_path}")

        doc = DocxTemplate(template_path)

        # Рендерим основные данные (титульный лист)
        doc.render(context)
        logger.info("✓ Шаблон успешно заполнен данными")

        # Применяем пост-обработчик для заполнения таблицы
        doc = _post_process_equipment_journal(doc, context)

        if equipment_type_name in LABEL_EQUIPMENT_TYPES:
            _append_equipment_labels(doc, equipment_list)
        if template_code == 'lestnicy-journal':
            _append_ladder_labels(doc, equipment_list, inspection_date=date_for_name)

        label = _get_equipment_label(equipment_type_name)
        date_for_name = _resolve_inspection_date(inspection_date) or start_date or datetime.date.today()
        date_str = date_for_name.strftime('%d.%m.%Y')
        org_name = _sanitize_filename(organization.short_name_ru)
        filename = f"Журнал осмотра {label} {org_name} {date_str}.docx"
        logger.info(f"Имя файла: {filename}")

        # Сохраняем в BytesIO
        docx_buffer = BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)

        file_content = docx_buffer.getvalue()
        logger.info(f"✓ Создан DOCX файл {filename}, размер: {len(file_content)} байт")

        return {
            'content': file_content,
            'filename': filename,
        }

    except Exception as e:
        logger.error(f"Ошибка при генерации журнала оборудования: {str(e)}")
        logger.error(traceback.format_exc())
        return None
