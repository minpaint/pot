"""
📄 Генератор для листа ознакомления (обновлённая версия 2)

Изменения v2 (после логов 17.04.2025):
1. Если в шаблоне нет табличного стиля «Table Grid», границы проставляются программно через XML.
2. Добавлена функция `_set_cell_border`, чтобы задать «single» 0.5 pt для всех сторон.
3. При установке стиля перехватываем `KeyError` и `ValueError`.
4. Код по‑прежнему задаёт Times New Roman 14 pt для всех ячеек.
"""

import logging
import traceback
import re
from typing import Dict, Any, Optional, List

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from directory.document_generators.base import (
    get_document_template,
    prepare_employee_context,
    generate_docx_from_template,
)

logger = logging.getLogger(__name__)


# =============================================
# Основная функция генерации листа ознакомления
# =============================================

def generate_familiarization_document(
    employee,
    document_list: Optional[List[str]] = None,
    user=None,
    custom_context: Optional[Dict[str, Any]] = None,
) -> Optional[Dict[str, Any]]:
    """Генерирует лист ознакомления для сотрудника."""

    try:
        template = get_document_template("doc_familiarization", employee)
        if not template:
            raise ValueError("Активный шаблон для листа ознакомления не найден")

        context = prepare_employee_context(employee)

        # Получаем список документов, если не передан вручную
        if document_list is None:
            from directory.views.documents.utils import get_employee_documents

            fetched_list, success = get_employee_documents(employee)
            document_list = (
                fetched_list
                if success
                else [
                    "Устав ЗАО 'СтройКомплекс'",
                    "Разрешение на строительство жилого комплекса 'Заречный'",
                    "Правила внутреннего трудового распорядка",
                ]
            )

        # Дополняем контекст
        context.update(
            {
                "documents_list": "DOCMARKER_START",
                "all_documents": document_list,
                "familiarization_date": context.get("current_date"),
            }
        )
        if custom_context:
            context.update(custom_context)

        # Рендер + пост‑обработка
        return generate_docx_from_template(
            template,
            context,
            employee,
            user,
            post_processor=process_table_rows,
        )

    except Exception as exc:
        logger.error("Ошибка при генерации листа ознакомления: %s", exc)
        logger.error(traceback.format_exc())
        return None


# =========================
#   ПОСТ‑ОБРАБОТКА ТАБЛИЦЫ
# =========================

def process_table_rows(doc, context):
    try:
        docs: List[str] = context.get("all_documents", [])
        if not docs:
            return doc

        docx_document = doc.docx

        for table in docx_document.tables:
            row_idx, doc_col = _find_template_row(table)
            if row_idx is None:
                continue

            date_col = _find_date_cell(table.rows[row_idx], doc_col)

            # Заполняем первую строку
            table.rows[row_idx].cells[doc_col].text = docs[0]

            # Добавляем новые строки
            for doc_name in docs[1:]:
                new_row = table.add_row()
                new_row.cells[doc_col].text = doc_name
                if date_col is not None and date_col < len(new_row.cells):
                    new_row.cells[date_col].text = context.get("familiarization_date", "")

            # Оформление
            _apply_table_format(table)
            break  # одна таблица – выходим

        return doc

    except Exception as exc:
        logger.error("Ошибка при пост‑обработке таблицы: %s", exc)
        logger.error(traceback.format_exc())
        return doc


# --------------------------
#   УТИЛИТЫ
# --------------------------

def _find_template_row(table):
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            if "DOCMARKER_START" in cell.text:
                return r_idx, c_idx
    return (1, 0) if len(table.rows) >= 2 else (None, None)


def _find_date_cell(row, doc_col):
    pattern = re.compile(r"\d{2}\.\d{2}\.\d{4}")
    for idx, cell in enumerate(row.cells):
        if idx == doc_col:
            continue
        if pattern.search(cell.text.strip()):
            return idx
    return len(row.cells) - 1 if len(row.cells) > doc_col + 1 else None


def _apply_table_format(table):
    """Times New Roman 14 pt, + все границы."""

    # 1. Попытка применить готовый стиль
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        logger.info("Стиль 'Table Grid' недоступен – проставляем границы вручную")
        _add_borders_manually(table)

    # 2. Шрифт всем run‑ам
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(14)


def _add_borders_manually(table):
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)


def _set_cell_border(
    cell,
    **kwargs,
):
    """Устанавливает одинарную границу толщиной 4 εм (≈0.5 pt) вокруг ячейки.
    Если передать свои параметры, можно переопределить любую сторону.
    """

    # Значения по умолчанию – все стороны single 4 εм
    sides = {
        "top": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "left": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "bottom": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "right": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
    }
    sides.update(kwargs)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tblBorders = tcPr.find(qn("w:tcBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tcBorders")
        tcPr.append(tblBorders)

    for edge in ("left", "top", "right", "bottom"):
        edge_el = tblBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tblBorders.append(edge_el)
        attrs = sides.get(edge, {})
        for key in ["val", "sz", "color", "space"]:
            edge_el.set(qn(f"w:{key}"), attrs.get(key, "single" if key == "val" else "4"))
