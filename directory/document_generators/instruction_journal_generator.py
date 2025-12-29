# directory/document_generators/instruction_journal_generator.py
"""
📄 Генератор для образца заполнения журнала повторных инструктажей
"""
import logging
import traceback
from typing import Dict, Any, Optional, List
from io import BytesIO
from docxtpl import DocxTemplate

from directory.document_generators.base import (
    get_document_template, prepare_employee_context
)
from directory.utils.vehicle_utils import combine_instructions
from directory.utils.declension import get_initials_from_name

# Настройка логирования
logger = logging.getLogger(__name__)


def _find_instruction_journal_table(docx_doc):
    """
    Находит таблицу журнала инструктажей в документе.
    """
    # Обычно это последняя (или единственная) таблица
    return docx_doc.tables[-1] if docx_doc.tables else None


def _set_cell_borders(cell):
    """
    Устанавливает границы для ячейки таблицы.
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.oxml.ns import qn

    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    # Удаляем старые границы, если есть
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None:
        tcPr.remove(tcBorders)

    # Создаём элемент границ
    tcBorders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
        </w:tcBorders>
    ''')

    # Добавляем новые границы
    tcPr.append(tcBorders)


def _reset_instruction_journal_table(table):
    """
    Удаляет все строки данных, оставляя только заголовки.
    Находит последнюю строку заголовка по первой ячейке (должна содержать "1").
    """
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    # Ищем последнюю строку заголовка - проверяем первую ячейку каждой строки
    last_header_row_idx = 1  # По умолчанию предполагаем, что это вторая строка (индекс 1)

    for row_idx, row in enumerate(table.rows):
        if len(row.cells) > 0:
            first_cell_text = row.cells[0].text.strip()
            logger.info(f"Проверка строки {row_idx}: первая ячейка = '{first_cell_text}'")
            # Если первая ячейка содержит "1" или "1." или "1)" - это строка с номерами столбцов
            # Проверяем разные варианты написания
            if (first_cell_text == '1' or
                first_cell_text == '1.' or
                first_cell_text == '1)' or
                first_cell_text.startswith('1') or
                '1' in first_cell_text[:3]):
                last_header_row_idx = row_idx
                logger.info(f"✓ Найдена строка с номерами столбцов: индекс {row_idx}, текст первой ячейки: '{first_cell_text}'")
                break

    # Количество строк заголовка = индекс строки с номерами + 1
    num_header_rows = last_header_row_idx + 1
    logger.info(f"Количество строк заголовка: {num_header_rows}")

    # Удаляем ВСЕ строки после заголовков
    while len(table.rows) > num_header_rows:
        row = table.rows[-1]
        tbl = table._tbl
        tbl.remove(row._tr)

    # Логируем количество столбцов в каждой строке
    for row_idx in range(num_header_rows):
        logger.info(f"Строка {row_idx}: количество ячеек = {len(table.rows[row_idx].cells)}")

    # Устанавливаем границы для всех ячеек в строках заголовка
    for row_idx in range(num_header_rows):
        for cell in table.rows[row_idx].cells:
            _set_cell_borders(cell)

    # ВАЖНО: Повторяем на каждой странице ВСЕ строки заголовка
    # Это более надежный способ - помечаем все строки от 0 до last_header_row_idx
    logger.info(f"Установка повторения для всех строк заголовка (0-{last_header_row_idx})")

    from docx.oxml.ns import qn

    for row_idx in range(num_header_rows):
        header_row = table.rows[row_idx]
        tr = header_row._tr
        trPr = tr.trPr

        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
            tr.insert(0, trPr)
            logger.info(f"  Создан новый trPr для строки {row_idx}")

        # Проверяем, нет ли уже тега tblHeader
        existing_tblHeader = trPr.find(qn('w:tblHeader'))
        if existing_tblHeader is not None:
            logger.info(f"  Тег w:tblHeader уже существует в строке {row_idx}, удаляем")
            trPr.remove(existing_tblHeader)

        # Добавляем тег tblHeader для повторения этой строки
        # ВАЖНО: w:tblHeader должен быть ПЕРВЫМ дочерним элементом в w:trPr
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trPr.insert(0, tblHeader)
        logger.info(f"  ✓ Установлено повторение для строки {row_idx}")

    logger.info(f"✓ Повторение установлено для всех {num_header_rows} строк заголовка")


def _fill_instruction_journal_rows(table, employees_data: List[Dict[str, str]], instruction_date: str, instruction_type: str, instruction_reason: str):
    """
    Заполняет таблицу журнала инструктажей строками с данными сотрудников.

    Args:
        table: Таблица документа Word
        employees_data: Данные сотрудников для заполнения
        instruction_date: Дата инструктажа
        instruction_type: Вид инструктажа
        instruction_reason: Причина проведения инструктажа
    """
    from docx.shared import Pt
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    for idx, emp in enumerate(employees_data, start=1):
        row = table.add_row()
        cells = row.cells
        cols = len(cells)

        # Логируем для первой строки данных
        if idx == 1:
            logger.info(f"Первая строка данных: количество ячеек = {cols}")

        # Заполняем ячейки данными
        # Структура журнала инструктажей (10 колонок):
        # 0: № п/п (оставляем пустым)
        # 1: Дата проведения инструктажа
        # 2: ФИО лица, прошедшего инструктаж
        # 3: Должность (профессия)
        # 4: Вид инструктажа
        # 5: Причина проведения (для внепланового/целевого)
        # 6: Номера инструкций
        # 7: ФИО проводившего инструктаж
        # 8: Подпись проводившего
        # 9: Подпись прошедшего

        if cols > 0:
            cells[0].text = ""  # Оставляем пустым для ручного заполнения
        if cols > 1:
            cells[1].text = instruction_date  # Дата
        if cols > 2:
            cells[2].text = emp.get('fio_initials', '')  # ФИО (Фамилия И.О.)
        if cols > 3:
            # Профессия/должность - с учётом подрядчиков
            if emp.get('is_contractor'):
                cells[3].text = emp.get('GPD', 'Работник по договору ГПХ')
            else:
                cells[3].text = emp.get('position_nominative', '')
        if cols > 4:
            cells[4].text = instruction_type  # Вид инструктажа
        if cols > 5:
            cells[5].text = instruction_reason  # Причина
        if cols > 6:
            cells[6].text = emp.get('instruction_numbers', '')  # Номера инструкций
        if cols > 7:
            # ФИО проводившего инструктаж
            cells[7].text = "Фамилия, инициалы руководителя"
        if cols > 8:
            # Подпись проводившего
            cells[8].text = "подпись работника"
        if cols > 9:
            # Подпись прошедшего
            cells[9].text = "подпись руководителя"
        if cols > 10:
            # Стажировка (оставляем пустым)
            cells[10].text = ""
        # Остальные ячейки (11+) оставляем пустыми для ручного заполнения (стажировка)

        # Запрещаем разрыв строки при переносе на новую страницу
        tr = row._tr
        trPr = tr.trPr
        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
            tr.insert(0, trPr)

        # Добавляем свойство cantSplit (не разрывать строку)
        cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
        trPr.append(cantSplit)

        # Применяем форматирование Times New Roman ко всем ячейкам строки
        for cell_idx, cell in enumerate(cells):
            _set_cell_borders(cell)  # Устанавливаем границы для каждой ячейки
            for paragraph in cell.paragraphs:
                # Центрируем: № п/п (0), Вид инструктажа (4), Подпись работника (8), Подпись руководителя (9)
                if cell_idx in [0, 4, 8, 9]:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    # Размеры шрифта по столбцам
                    if cell_idx == 7:  # Столбец 8: ФИО проводившего - 10 кегль
                        run.font.size = Pt(10)
                    elif cell_idx in [8, 9]:  # Столбцы 9-10: подписи - 9 кегль
                        run.font.size = Pt(9)
                    else:  # Остальные столбцы - 12 кегль
                        run.font.size = Pt(12)


def generate_instruction_journal(
    employees: List,
    date_povtorny: str,
    user=None,
    custom_context: Optional[Dict[str, Any]] = None,
    grouping_name: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """
    Генерирует образец заполнения журнала повторных инструктажей для списка сотрудников.

    Args:
        employees: Список объектов Employee
        date_povtorny: Дата повторного инструктажа в формате строки
        user: Пользователь, создающий документ (опционально)
        custom_context: Пользовательский контекст (опционально)
        grouping_name: Название группы (подразделения) для имени файла

    Returns:
        Optional[Dict]: Словарь с 'content' и 'filename' или None при ошибке
    """
    try:
        logger.info(f"Начало генерации образца журнала инструктажей для {len(employees)} сотрудников")

        if not employees:
            logger.error("Не переданы сотрудники для образца журнала")
            raise ValueError("Не переданы сотрудники для образца журнала")

        # Получаем первого сотрудника для определения организации
        first_employee = employees[0]
        logger.info(f"Первый сотрудник: {first_employee.full_name_nominative}")

        # Получаем шаблон
        logger.info("Поиск шаблона типа 'instruction_journal'")
        template = get_document_template('instruction_journal', first_employee)
        if not template:
            logger.error("Активный шаблон для образца журнала инструктажей не найден")
            raise ValueError("Активный шаблон для образца журнала инструктажей не найден. Создайте шаблон в админке.")

        # Подготавливаем контекст для первого сотрудника (для общей информации)
        context = prepare_employee_context(first_employee)

        # Добавляем дату инструктажа в формате ДД.ММ.ГГГГ
        from datetime import datetime
        try:
            # Преобразуем дату из формата YYYY-MM-DD в DD.MM.YYYY
            date_obj = datetime.strptime(date_povtorny, '%Y-%m-%d')
            formatted_date = date_obj.strftime('%d.%m.%Y')
            context['instruction_date'] = formatted_date
        except:
            # Если не удалось преобразовать, используем как есть
            context['instruction_date'] = date_povtorny
            formatted_date = date_povtorny

        # Добавляем вид инструктажа и причину (по умолчанию)
        context['instruction_type'] = 'Повторный'
        context['instruction_reason'] = ''

        # Добавляем название группы, если указано
        if grouping_name:
            context['grouping_name'] = grouping_name

        # Добавляем иерархический заголовок: отдел → подразделение → организация
        # Выбираем наиболее конкретный уровень структуры
        if first_employee.department:
            context['structural_unit'] = first_employee.department.name
            context['structural_unit_genitive'] = context.get('department_genitive', '')
            context['structural_unit_dative'] = context.get('department_dative', '')
        elif first_employee.subdivision:
            context['structural_unit'] = first_employee.subdivision.name
            context['structural_unit_genitive'] = context.get('subdivision_genitive', '')
            context['structural_unit_dative'] = context.get('subdivision_dative', '')
        elif first_employee.organization:
            context['structural_unit'] = first_employee.organization.short_name_ru
            context['structural_unit_genitive'] = context.get('organization_name_genitive', '')
            context['structural_unit_dative'] = context.get('organization_name_dative', '')
        else:
            context['structural_unit'] = ''
            context['structural_unit_genitive'] = ''
            context['structural_unit_dative'] = ''

        # Добавляем пользовательский контекст (переопределяет значения по умолчанию)
        if custom_context:
            context.update(custom_context)

        # Получаем финальные значения для использования в таблице
        instruction_type = context.get('instruction_type', 'Повторный')
        instruction_reason = context.get('instruction_reason', '')

        # Подготавливаем список сотрудников для заполнения таблицы
        employees_data = []
        for emp in employees:
            # Получаем все инструкции для сотрудника
            instruction_numbers = combine_instructions(emp)

            # Определяем, является ли сотрудник подрядчиком
            is_contractor = getattr(emp, 'contract_type', 'standard') == 'contractor'

            employee_data = {
                'fio_nominative': emp.full_name_nominative or '',
                'fio_initials': get_initials_from_name(emp.full_name_nominative or ''),
                'position_nominative': emp.position.position_name if emp.position else '',
                'instruction_numbers': instruction_numbers,
                'is_contractor': is_contractor,
                'GPD': 'Работник по договору ГПХ' if is_contractor else '',
            }
            employees_data.append(employee_data)

        logger.info(f"Подготовлено {len(employees_data)} сотрудников для заполнения таблицы")

        # Загружаем и рендерим шаблон с помощью DocxTemplate
        doc = DocxTemplate(template.template_file.path)

        # Рендерим переменные шаблона (заголовки и т.д.)
        render_context = context.copy()
        render_context.pop('employee', None)  # Удаляем объект employee из контекста
        logger.info("Рендеринг шаблона с контекстом")
        doc.render(render_context)

        # Находим и заполняем таблицу журнала инструктажей
        table = _find_instruction_journal_table(doc.docx)
        if table:
            logger.info(f"ПОСЛЕ рендеринга: таблица содержит {len(table.rows)} строк")
            logger.info("Таблица журнала найдена, очищаем и заполняем данными")
            _reset_instruction_journal_table(table)
            logger.info(f"ПОСЛЕ сброса: таблица содержит {len(table.rows)} строк")
            _fill_instruction_journal_rows(table, employees_data, formatted_date, instruction_type, instruction_reason)
            logger.info(f"ПОСЛЕ заполнения: таблица содержит {len(table.rows)} строк")
        else:
            logger.warning("Таблица журнала не найдена в шаблоне")

        # Сохраняем финальный документ в буфер
        final_buffer = BytesIO()
        doc.save(final_buffer)
        final_buffer.seek(0)

        # Формируем имя файла по иерархии: отдел → подразделение → организация
        if grouping_name:
            # Очищаем название от недопустимых символов для имени файла
            safe_name = grouping_name.replace('"', '').replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            filename = f"Образец_журнала_инструктажей_{safe_name}.docx"
        else:
            # Используем ту же иерархическую логику, что и для заголовков
            if first_employee.department:
                unit_name = first_employee.department.name
            elif first_employee.subdivision:
                unit_name = first_employee.subdivision.name
            elif first_employee.organization:
                unit_name = first_employee.organization.short_name_ru
            else:
                # Если нет структурной единицы, используем инициалы первого сотрудника
                unit_name = get_initials_from_name(first_employee.full_name_nominative)

            # Очищаем название от недопустимых символов для имени файла
            safe_name = unit_name.replace('"', '').replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('<', '_').replace('>', '_').replace('|', '_')
            filename = f"Образец_журнала_инструктажей_{safe_name}.docx"

        logger.info(f"Образец журнала инструктажей успешно сгенерирован: {filename}")
        return {'content': final_buffer.getvalue(), 'filename': filename}

    except Exception as e:
        logger.error(f"Ошибка при генерации образца журнала инструктажей: {str(e)}")
        logger.error(traceback.format_exc())
        return None
