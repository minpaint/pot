import logging
import datetime
import traceback
from io import BytesIO
from pathlib import Path
from typing import Dict, Any, Optional, List

from docxtpl import DocxTemplate
from django.conf import settings

from directory.document_generators.base import (
    get_document_template,
    prepare_employee_context,
    generate_docx_from_template,
)

# Сервисные функции для работы с комиссией (экспортируемые из directory/utils/__init__.py)
from directory.utils import find_appropriate_commission, get_commission_members_formatted
# Для склонения названий в родительном падеже
from directory.utils.declension import decline_phrase
# Для очистки пустых параграфов в сгенерированных документах
from directory.utils.docx_cleaner import clean_document

logger = logging.getLogger(__name__)


def generate_knowledge_protocol(
    employee,
    user=None,
    custom_context: Optional[Dict[str, Any]] = None
) -> Optional[Dict[str, Any]]:
    """
    Генерирует протокол проверки знаний по вопросам охраны труда для сотрудника.

    Args:
        employee: Объект модели Employee
        user: Пользователь, создающий документ (опционально)
        custom_context: Пользовательский контекст (опционально)

    Returns:
        Optional[Dict]: Словарь с 'content' и 'filename' или None при ошибке
    """
    try:
        # 1) Шаблон
        template = get_document_template('knowledge_protocol', employee)
        if not template:
            raise ValueError("Не найден активный шаблон 'knowledge_protocol'")

        # 2) Базовый контекст
        context = prepare_employee_context(employee)
        # Базовый контекст уже содержит правильные значения для должности/работы по договору подряда
        # в полях position_nominative, position_genitive и т.д.

        # 3) Номер и дата протокола
        now = datetime.datetime.now()
        context.setdefault('protocol_number', f"PZ-{now.strftime('%Y%m%d')}-{employee.id}")
        context.setdefault('protocol_date', now.strftime("%d.%m.%Y"))

        # 4) Комиссия и её состав
        commission = find_appropriate_commission(employee)
        cdata = get_commission_members_formatted(commission) if commission else {}

        # 4.1) Председатель
        chairman = cdata.get('chairman', {})
        has_chairman = bool(chairman.get('name'))
        context.setdefault('chairman_role_label', 'Председатель комиссии:' if has_chairman else '')
        context.setdefault('chairman_name', chairman.get('name', ''))
        context.setdefault('chairman_position', chairman.get('position', '').lower() if chairman.get('position') else '')
        context.setdefault('chairman_name_initials', chairman.get('name_initials', ''))

        # 4.2) Заместитель председателя
        vice_chairman = cdata.get('vice_chairman', {})
        has_vice_chairman = bool(vice_chairman.get('name'))
        context.setdefault('vice_chairman_role_label', 'Заместитель председателя комиссии:' if has_vice_chairman else '')
        context.setdefault('vice_chairman_name', vice_chairman.get('name', ''))
        context.setdefault('vice_chairman_position', vice_chairman.get('position', '').lower() if vice_chairman.get('position') else '')
        context.setdefault('vice_chairman_name_initials', vice_chairman.get('name_initials', ''))

        # 4.3) Секретарь
        secretary = cdata.get('secretary', {})
        has_secretary = bool(secretary.get('name'))
        context.setdefault('secretary_role_label', 'Секретарь комиссии:' if has_secretary else '')
        context.setdefault('secretary_name', secretary.get('name', ''))
        context.setdefault('secretary_position', secretary.get('position', '').lower() if secretary.get('position') else '')
        context.setdefault('secretary_name_initials', secretary.get('name_initials', ''))

        # 4.4) Члены комиссии
        members = cdata.get('members_formatted', [])
        context.setdefault('members_formatted', members)

        # 4.5) Параграфы «ФИО – должность»
        members_paragraphs = [
            f"{m['name']} - {m['position'].lower()}"
            for m in members
        ]
        context['members_paragraphs'] = members_paragraphs

        # 4.6) Параграфы с инициалами
        members_initials_paragraphs = [
            m['name_initials'] for m in members
        ]
        context['members_initials_paragraphs'] = members_initials_paragraphs

        # 4.6) binding_name_genitive — «протокол комиссии чего…»
        if commission:
            if commission.department:
                binding = decline_phrase(commission.department.name, 'gent')
            elif commission.subdivision:
                binding = decline_phrase(commission.subdivision.name, 'gent')
            elif commission.organization:
                # Название организации НЕ склоняется - это имя собственное в кавычках
                # "Комиссия ООО "Безопасность Плюс"", а не "ооо безопасности Плюс"
                binding = commission.organization.short_name_ru
            else:
                binding = ""
        else:
            binding = ""
        context.setdefault('binding_name_genitive', binding)

        # 5) Подмешать custom_context, если есть
        if custom_context:
            context.update(custom_context)

        logger.debug(f"[generate_knowledge_protocol] context keys: {list(context.keys())}")

        # 6) Рендерим шаблон с помощью docxtpl
        from docxtpl import DocxTemplate
        from pathlib import Path

        template_path = template.template_file.path
        doc = DocxTemplate(template_path)

        render_context = context.copy()
        render_context.pop('employee', None)
        doc.render(render_context)

        # 7) Заполняем таблицу результатов проверки знаний
        from directory.utils.vehicle_utils import needs_vehicle_training, get_vehicle_position_name

        table = _find_knowledge_protocol_table(doc.docx)
        if table:
            # Очищаем все строки кроме заголовка
            _reset_periodic_table(table)

            # Формируем данные для строк
            employees_data = []

            # Строка 1: проверка знаний по профессии (основная должность)
            employees_data.append({
                'fio_nominative': context.get('fio_nominative', ''),
                'position_nominative': context.get('position_nominative', ''),
                'ticket_number': custom_context.get('ticket_number', '') if custom_context else '',
            })

            # Строка 2: проверка знаний по видам выполняемых работ (управление автомобилем)
            if needs_vehicle_training(employee):
                employees_data.append({
                    'fio_nominative': context.get('fio_nominative', ''),
                    'position_nominative': get_vehicle_position_name(),
                    'ticket_number': custom_context.get('ticket_number', '') if custom_context else '',
                })

            # Строки 3+: проверка знаний по видам ответственности
            if employee.position and employee.position.responsibility_types.exists():
                for resp_type in employee.position.responsibility_types.filter(is_active=True).order_by('order', 'name'):
                    employees_data.append({
                        'fio_nominative': context.get('fio_nominative', ''),
                        'position_nominative': resp_type.name,
                        'ticket_number': custom_context.get('ticket_number', '') if custom_context else '',
                    })

            # Заполняем таблицу с типом проверки "первичная" (при приёме на работу)
            _fill_periodic_rows(table, employees_data, check_type='первичная')

            # Удаляем лишние пустые параграфы после таблицы
            _remove_empty_paragraphs_after_table(doc.docx, table)

        # 8) Сохраняем документ
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # 9) Очищаем пустые параграфы и строки
        logger.info("[generate_protocol] ПЕРЕД вызовом clean_document")
        cleaned_content = clean_document(buffer.getvalue())
        logger.info("[generate_protocol] ПОСЛЕ вызова clean_document")

        from directory.utils.declension import get_initials_from_name
        employee_initials = get_initials_from_name(context.get('fio_nominative', ''))
        filename = f"Протокол_{employee_initials}.docx"

        return {'content': cleaned_content, 'filename': filename}

    except Exception:
        logger.error("Ошибка генерации протокола", exc_info=True)
        return None


def _find_knowledge_protocol_table(docx_doc):
    """
    Находит таблицу результатов в первичном протоколе проверки знаний.
    Ищет таблицу с заголовком содержащим "Фамилия" и "должность".
    """
    for table in docx_doc.tables:
        if not table.rows:
            continue
        header_text = " ".join(cell.text for cell in table.rows[0].cells)
        if "Фамилия" in header_text and "должность" in header_text:
            return table
    # Если не нашли по ключевым словам, берём первую таблицу
    return docx_doc.tables[0] if docx_doc.tables else None


def _find_periodic_table(docx_doc):
    """
    Locate the protocol results table by header keywords.
    """
    for table in docx_doc.tables:
        if not table.rows:
            continue
        header_text = " ".join(cell.text for cell in table.rows[0].cells)
        if "Результаты проверки знаний" in header_text and "Роспись" in header_text:
            return table
    return docx_doc.tables[-1] if docx_doc.tables else None


def _reset_periodic_table(table):
    """Remove all data rows keeping the header row only."""
    while len(table.rows) > 1:
        row = table.rows[-1]
        tbl = table._tbl
        tbl.remove(row._tr)

    # Устанавливаем повторение шапки таблицы на каждой странице
    if table.rows:
        from docx.oxml import parse_xml
        from docx.oxml.ns import nsdecls

        header_row = table.rows[0]
        # Получаем или создаем элемент trPr (table row properties)
        tr = header_row._tr
        trPr = tr.trPr
        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
            tr.insert(0, trPr)

        # Добавляем тег tblHeader для повторения строки
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")}/>')
        trPr.append(tblHeader)


def _remove_empty_paragraphs_after_table(doc, table):
    """
    Удаляет пустые параграфы после таблицы, чтобы избежать пустых страниц.
    """
    # Получаем индекс таблицы в документе
    table_element = table._element
    body = doc.element.body

    # Находим индекс таблицы
    table_index = None
    for idx, element in enumerate(body):
        if element == table_element:
            table_index = idx
            break

    if table_index is None:
        return

    # Удаляем все параграфы после таблицы, которые пустые или содержат только пробелы
    elements_to_remove = []
    for idx in range(table_index + 1, len(body)):
        element = body[idx]
        # Проверяем только параграфы
        if element.tag.endswith('}p'):
            # Получаем текст параграфа
            text = ''.join(node.text for node in element.iter() if node.text)
            # Если параграф пустой или содержит только пробелы - помечаем на удаление
            if not text or text.strip() == '':
                elements_to_remove.append(element)
            else:
                # Если нашли непустой параграф - прекращаем проверку
                break

    # Удаляем помеченные элементы
    for element in elements_to_remove:
        body.remove(element)


def _fill_periodic_rows(table, employees_data: List[Dict[str, str]], check_type: str = 'периодическая'):
    """
    Append rows with employee data to the protocol table.

    Args:
        table: Таблица документа Word
        employees_data: Данные сотрудников для заполнения
        check_type: Тип проверки знаний ('первичная' или 'периодическая')
    """
    from docx.shared import Pt
    from docx.oxml import parse_xml
    from docx.oxml.ns import nsdecls

    for idx, emp in enumerate(employees_data, start=1):
        row = table.add_row()
        cells = row.cells
        cols = len(cells)

        # Заполняем ячейки данными
        # Используем сквозную нумерацию (стандартная практика для протоколов)
        cells[0].text = str(idx)
        if cols > 1:
            cells[1].text = emp.get('fio_nominative', '')
        if cols > 2:
            cells[2].text = emp.get('position_nominative', '')
        if cols > 3:
            cells[3].text = check_type
        if cols > 4:
            ticket = emp.get('ticket_number') or ""
            cells[4].text = str(ticket)
        if cols > 5:
            cells[5].text = ""
        if cols > 6:
            cells[6].text = ""

        # Запрещаем разрыв строки при переносе на новую страницу
        tr = row._tr
        trPr = tr.trPr
        if trPr is None:
            trPr = parse_xml(f'<w:trPr {nsdecls("w")}/>')
            tr.insert(0, trPr)

        # Добавляем свойство cantSplit (не разрывать строку)
        cantSplit = parse_xml(f'<w:cantSplit {nsdecls("w")}/>')
        trPr.append(cantSplit)

        # Применяем форматирование Times New Roman ко всем ячейкам строки
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        for cell_idx, cell in enumerate(cells):
            for paragraph in cell.paragraphs:
                # Центрируем первую колонку (№ п/п)
                if cell_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)


def generate_periodic_protocol(
    employees: List,
    user=None,
    custom_context: Optional[Dict[str, Any]] = None,
    grouping_name: Optional[str] = None,
) -> Optional[Dict[str, Any]]:
    """
    Сформировать протокол периодической проверки знаний для списка сотрудников.
    """
    print(f"\n\n===== DEBUG: ВЫЗОВ generate_periodic_protocol для {len(employees) if employees else 0} сотрудников =====\n", flush=True)
    logger.info(f"===== ВЫЗОВ generate_periodic_protocol для {len(employees) if employees else 0} сотрудников =====")
    try:
        if not employees:
            raise ValueError("Не переданы сотрудники для протокола")

        primary_employee = employees[0]
        template = get_document_template('periodic_protocol', primary_employee)
        fallback_path = Path(settings.MEDIA_ROOT) / 'document_templates' / 'etalon' / 'periodic_protocol_template.docx'
        if template:
            template_path = template.template_file.path
        elif fallback_path.exists():
            template_path = str(fallback_path)
        else:
            raise ValueError("Не найден активный шаблон 'periodic_protocol'")

        context = prepare_employee_context(primary_employee)
        now = datetime.datetime.now()
        # Протокол номер/дата оставляем пустыми по требованию
        context.setdefault('protocol_number', "")
        context.setdefault('protocol_date', "")

        commission = find_appropriate_commission(primary_employee)
        logger.info(f"[periodic_protocol] Найдена комиссия: {commission}")
        logger.info(f"[periodic_protocol] Сотрудник: {primary_employee.full_name_nominative}, org={primary_employee.organization_id}, subdiv={primary_employee.subdivision_id}, dept={primary_employee.department_id}")

        cdata = get_commission_members_formatted(commission) if commission else {}
        logger.info(f"[periodic_protocol] Данные комиссии: chairman={bool(cdata.get('chairman'))}, secretary={bool(cdata.get('secretary'))}, members={len(cdata.get('members_formatted', []))}")

        chairman = cdata.get('chairman', {})
        has_chairman = bool(chairman.get('name'))
        context.setdefault('chairman_role_label', 'Председатель комиссии:' if has_chairman else '')
        context.setdefault('chairman_name', chairman.get('name', ''))
        context.setdefault('chairman_position', chairman.get('position', '').lower() if chairman.get('position') else '')
        context.setdefault('chairman_name_initials', chairman.get('name_initials', ''))

        vice_chairman = cdata.get('vice_chairman', {})
        has_vice_chairman = bool(vice_chairman.get('name'))
        context.setdefault('vice_chairman_role_label', 'Заместитель председателя комиссии:' if has_vice_chairman else '')
        context.setdefault('vice_chairman_name', vice_chairman.get('name', ''))
        context.setdefault('vice_chairman_position', vice_chairman.get('position', '').lower() if vice_chairman.get('position') else '')
        context.setdefault('vice_chairman_name_initials', vice_chairman.get('name_initials', ''))

        secretary = cdata.get('secretary', {})
        has_secretary = bool(secretary.get('name'))
        context.setdefault('secretary_role_label', 'Секретарь комиссии:' if has_secretary else '')
        context.setdefault('secretary_name', secretary.get('name', ''))
        context.setdefault('secretary_position', secretary.get('position', '').lower() if secretary.get('position') else '')
        context.setdefault('secretary_name_initials', secretary.get('name_initials', ''))

        members = cdata.get('members_formatted', [])
        context.setdefault('members_formatted', members)

        # Форматирование членов комиссии для шаблона
        members_paragraphs = [
            f"{m['name']} - {m['position'].lower()}"
            for m in members
        ]
        context['members_paragraphs'] = members_paragraphs
        logger.info(f"[periodic_protocol] Сформировано {len(members_paragraphs)} членов комиссии: {members_paragraphs}")

        # Параграфы с инициалами для членов комиссии
        members_initials_paragraphs = [
            m['name_initials'] for m in members
        ]
        context['members_initials_paragraphs'] = members_initials_paragraphs

        # Определяем binding по уровню найденной комиссии, а не по grouping_name
        # grouping_name используется только для группировки файлов, но не влияет на состав комиссии
        if commission:
            if commission.department:
                binding = decline_phrase(commission.department.name, 'gent')
            elif commission.subdivision:
                binding = decline_phrase(commission.subdivision.name, 'gent')
            elif commission.organization:
                binding = commission.organization.short_name_ru
            else:
                binding = ""
        else:
            binding = context.get('organization_name_genitive', "")
        context.setdefault('binding_name_genitive', binding)

        if custom_context:
            context.update(custom_context)

        # Импортируем утилиты для управления автомобилем
        from directory.utils.vehicle_utils import (
            needs_vehicle_training,
            get_vehicle_position_name,
        )

        employees_data = []
        for idx, emp in enumerate(employees, start=1):
            emp_ctx = prepare_employee_context(emp)

            # Проверка знаний по профессии - всегда добавляем основную должность
            employees_data.append({
                'fio_nominative': emp_ctx.get('fio_nominative', ''),
                'position_nominative': emp_ctx.get('position_nominative', ''),
                'ticket_number': '',  # Номер билета оставляем пустым для ручного заполнения
            })

            # Проверка знаний по видам выполняемых работ
            # Если сотрудник управляет автомобилем - добавляем вторую строку
            if needs_vehicle_training(emp):
                employees_data.append({
                    'fio_nominative': emp_ctx.get('fio_nominative', ''),
                    'position_nominative': get_vehicle_position_name(),
                    'ticket_number': '',
                })

            # Проверка знаний по видам ответственности
            if emp.position and emp.position.responsibility_types.exists():
                for resp_type in emp.position.responsibility_types.filter(is_active=True).order_by('order', 'name'):
                    employees_data.append({
                        'fio_nominative': emp_ctx.get('fio_nominative', ''),
                        'position_nominative': resp_type.name,
                        'ticket_number': '',
                    })

        doc = DocxTemplate(template_path)
        render_context = context.copy()
        render_context.pop('employee', None)
        doc.render(render_context)

        table = _find_periodic_table(doc.docx)
        if table:
            _reset_periodic_table(table)
            # Заполняем таблицу с типом проверки "периодическая"
            _fill_periodic_rows(table, employees_data, check_type='периодическая')

            # Удаляем лишние пустые параграфы после таблицы
            _remove_empty_paragraphs_after_table(doc.docx, table)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        # Очищаем пустые параграфы и строки
        print(f"\n===== DEBUG: ПЕРЕД вызовом clean_document, размер документа: {len(buffer.getvalue())} байт =====\n", flush=True)
        logger.info("[generate_periodic_protocols] ПЕРЕД вызовом clean_document")
        cleaned_content = clean_document(buffer.getvalue())
        print(f"\n===== DEBUG: ПОСЛЕ вызова clean_document, размер: {len(cleaned_content)} байт =====\n", flush=True)
        logger.info("[generate_periodic_protocols] ПОСЛЕ вызова clean_document")

        # Формируем имя файла на основе grouping_name или организации
        if grouping_name:
            # Если передано название подразделения для группировки
            # Убираем кавычки из названия файла
            clean_name = grouping_name.replace('"', '').replace("'", '').replace('«', '').replace('»', '')
            filename = f"Протокол проверки знаний по ОТ_{clean_name}.docx"
        else:
            # Используем название организации первого сотрудника
            org_name = primary_employee.organization.short_name_ru if primary_employee.organization else "Организация"
            # Убираем кавычки из названия файла
            clean_name = org_name.replace('"', '').replace("'", '').replace('«', '').replace('»', '')
            filename = f"Протокол проверки знаний по ОТ_{clean_name}.docx"

        return {'content': cleaned_content, 'filename': filename}

    except Exception:
        logger.error("Ошибка формирования периодического протокола", exc_info=True)
        return None
