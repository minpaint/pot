# directory/document_generators/siz_card_docx_generator.py
"""
üõ°Ô∏è –ì–µ–Ω–µ—Ä–∞—Ç–æ—Ä –ª–∏—á–Ω–æ–π –∫–∞—Ä—Ç–æ—á–∫–∏ —É—á—ë—Ç–∞ –°–ò–ó (DOCX).

- –ò—Å–ø–æ–ª—å–∑—É–µ—Ç –º–µ—Ö–∞–Ω–∏–∑–º –º–∞—Ä–∫–µ—Ä–æ–≤ –∞–Ω–∞–ª–æ–≥–∏—á–Ω—ã–π –ª–∏—Å—Ç—É –æ–∑–Ω–∞–∫–æ–º–ª–µ–Ω–∏—è
- –ü–æ–¥–¥–µ—Ä–∂–∏–≤–∞–µ—Ç –æ–±—ä–µ–¥–∏–Ω–µ–Ω–Ω—ã–µ —Å—Ç—Ä–æ–∫–∏ –¥–ª—è —É—Å–ª–æ–≤–∏–π –Ω–∞ –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω–µ
- –ó–∞–ø–æ–ª–Ω—è–µ—Ç —Ç–æ–ª—å–∫–æ –Ω–µ–æ–±—Ö–æ–¥–∏–º—ã–µ –∫–æ–ª–æ–Ω–∫–∏ –Ω–∞ –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω–µ
- –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ—Ç –∑–∞–≥–æ–ª–æ–≤–∫–∏ —Ç–∞–±–ª–∏—Ü —à—Ä–∏—Ñ—Ç–æ–º 9 –ø—É–Ω–∫—Ç–æ–≤
- –ü—Ä–∏–º–µ–Ω—è–µ—Ç –µ–¥–∏–Ω—ã–π —Å—Ç–∏–ª—å —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏—è –¥–ª—è –≤—Å–µ—Ö —Å—Ç—Ä–æ–∫ —Ç–∞–±–ª–∏—Ü—ã
- –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç –æ–¥–∏–Ω–∞—Ä–Ω—ã–π –º–µ–∂–¥—É—Å—Ç—Ä–æ—á–Ω—ã–π –∏–Ω—Ç–µ—Ä–≤–∞–ª
- –£–±–∏—Ä–∞–µ—Ç –æ—Ç—Å—Ç—É–ø—ã –¥–æ –∏ –ø–æ—Å–ª–µ –∞–±–∑–∞—Ü–µ–≤
- –ì–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç —Å–ª—É—á–∞–π–Ω—ã–µ —Ä–∞–∑–º–µ—Ä—ã –°–ò–ó –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç –ø–æ–ª–∞
"""

import logging
import traceback
import re
import random
from typing import Dict, Any, Optional, List, Tuple

from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

from directory.document_generators.base import (
    get_document_template,
    prepare_employee_context,
    generate_docx_from_template,
)
from directory.models.siz import SIZNorm
from directory.models.siz_issued import SIZIssued
from directory.utils.siz_sizes import get_employee_sizes

logger = logging.getLogger(__name__)


# =============================================
# –û—Å–Ω–æ–≤–Ω–∞—è —Ñ—É–Ω–∫—Ü–∏—è –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–∞—Ä—Ç–æ—á–∫–∏ –°–ò–ó
# =============================================

def generate_siz_card_docx(
        employee,
        user=None,
        custom_context: Optional[Dict[str, Any]] = None,
        raise_on_error: bool = False,
) -> Optional[Dict[str, Any]]:
    """–ì–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç –∫–∞—Ä—Ç–æ—á–∫—É —É—á—ë—Ç–∞ –°–ò–ó –¥–ª—è —Å–æ—Ç—Ä—É–¥–Ω–∏–∫–∞."""

    try:
        # 1. –ü–æ–ª—É—á–µ–Ω–∏–µ —à–∞–±–ª–æ–Ω–∞
        template = get_document_template("siz_card", employee)
        if not template:
            raise ValueError("–ê–∫—Ç–∏–≤–Ω—ã–π —à–∞–±–ª–æ–Ω –¥–ª—è –∫–∞—Ä—Ç–æ—á–∫–∏ –°–ò–ó –Ω–µ –Ω–∞–π–¥–µ–Ω")

        # 2. –ü–æ–¥–≥–æ—Ç–æ–≤–∫–∞ –±–∞–∑–æ–≤–æ–≥–æ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞
        context = prepare_employee_context(employee)
        full_name = context.get("fio_nominative", "")

        # –†–∞–∑–¥–µ–ª—è–µ–º –ø–æ–ª–Ω–æ–µ –∏–º—è –Ω–∞ —Å–æ—Å—Ç–∞–≤–ª—è—é—â–∏–µ
        last_name, first_name, patronymic = _split_full_name(full_name)

        # 3. –û–ø—Ä–µ–¥–µ–ª–µ–Ω–∏–µ –ø–æ–ª–∞ –¥–ª—è –∑–∞–≥–æ–ª–æ–≤–∫–∞
        gender = _gender_from_patronymic(patronymic)

        # 4. –ì–µ–Ω–µ—Ä–∞—Ü–∏—è —Å–ª—É—á–∞–π–Ω—ã—Ö —Ä–∞–∑–º–µ—Ä–æ–≤ –°–ò–ó
        ppe_head, ppe_gloves, sizod, _ = _generate_random_ppe_sizes(gender)

        # 5. –ü–æ–ª—É—á–∞–µ–º –≤—ã–±—Ä–∞–Ω–Ω—ã–µ –Ω–æ—Ä–º—ã –°–ò–ó –∏–∑ GET-–ø–∞—Ä–∞–º–µ—Ç—Ä–æ–≤ (–¥–ª—è –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã)
        selected_norm_ids = []
        if custom_context:
            if 'selected_norm_ids' in custom_context:
                selected_norm_ids = custom_context['selected_norm_ids']
            elif 'selected_norms' in custom_context:
                selected_norm_ids = custom_context['selected_norms']

        # 6. –ü–æ–ª—É—á–∞–µ–º –í–°–ï –Ω–æ—Ä–º—ã –°–ò–ó –¥–ª—è –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω—ã (–Ω–µ–∑–∞–≤–∏—Å–∏–º–æ –æ—Ç –≤—ã–±–æ—Ä–∞)
        all_norms_data = []
        if employee.position:
            # –°–Ω–∞—á–∞–ª–∞ –ø—ã—Ç–∞–µ–º—Å—è –ø–æ–ª—É—á–∏—Ç—å –Ω–æ—Ä–º—ã –¥–ª—è –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ–π –¥–æ–ª–∂–Ω–æ—Å—Ç–∏
            all_norms_query = SIZNorm.objects.filter(
                position=employee.position
            ).select_related('siz')

            # –ï—Å–ª–∏ –Ω–æ—Ä–º—ã –Ω–µ –Ω–∞–π–¥–µ–Ω—ã, –∏—â–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—É—é –¥–æ–ª–∂–Ω–æ—Å—Ç—å
            if not all_norms_query.exists():
                logger.info("–ù–æ—Ä–º—ã –Ω–µ –Ω–∞–π–¥–µ–Ω—ã –¥–ª—è –∫–æ–Ω–∫—Ä–µ—Ç–Ω–æ–π –¥–æ–ª–∂–Ω–æ—Å—Ç–∏, –∏—â–µ–º —ç—Ç–∞–ª–æ–Ω–Ω—É—é...")
                from directory.models import Position

                positions_with_same_name = Position.objects.filter(
                    position_name=employee.position.position_name
                ).order_by('organization__full_name_ru')

                reference_position = None
                for pos in positions_with_same_name:
                    if SIZNorm.objects.filter(position=pos).exists():
                        reference_position = pos
                        break

                if reference_position:
                    logger.info(f"–ù–∞–π–¥–µ–Ω–∞ —ç—Ç–∞–ª–æ–Ω–Ω–∞—è –¥–æ–ª–∂–Ω–æ—Å—Ç—å ID={reference_position.id}")
                    all_norms_query = SIZNorm.objects.filter(
                        position=reference_position
                    ).select_related('siz')

            for norm in all_norms_query:
                cost = norm.siz.cost
                cost_display = f"{cost:.2f}" if cost is not None else ""
                all_norms_data.append({
                    "name": norm.siz.name,
                    "classification": norm.siz.classification,
                    "unit": norm.siz.unit,
                    "quantity": norm.quantity,
                    "wear_period": "–î–æ –∏–∑–Ω–æ—Å–∞" if norm.siz.wear_period == 0 else str(norm.siz.wear_period),
                    "condition": norm.condition,
                    "cost": cost_display,
                    "id": norm.id  # –î–æ–±–∞–≤–ª—è–µ–º ID –¥–ª—è –¥–∞–ª—å–Ω–µ–π—à–µ–≥–æ –∏—Å–ø–æ–ª—å–∑–æ–≤–∞–Ω–∏—è
                })

        # 7. –ü–æ–ª—É—á–∞–µ–º –í–´–ë–†–ê–ù–ù–´–ï –Ω–æ—Ä–º—ã –°–ò–ó –¥–ª—è –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã
        selected_norms_data = []
        if selected_norm_ids:
            # –§–∏–ª—å—Ç—Ä—É–µ–º –ø–æ–ª–Ω—ã–π —Å–ø–∏—Å–æ–∫ –ø–æ –≤—ã–±—Ä–∞–Ω–Ω—ã–º ID
            selected_norms_data = [norm for norm in all_norms_data if str(norm["id"]) in selected_norm_ids]
        else:
            # –ï—Å–ª–∏ ID –Ω–µ –≤—ã–±—Ä–∞–Ω—ã, –∏—Å–ø–æ–ª—å–∑—É–µ–º –≤—Å–µ –Ω–æ—Ä–º—ã (–ø–æ —É–º–æ–ª—á–∞–Ω–∏—é)
            selected_norms_data = all_norms_data.copy()

        # 8. –§–æ—Ä–º–∏—Ä–æ–≤–∞–Ω–∏–µ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞ —Å –∑–∞–≥–ª—É—à–∫–∞–º–∏ –¥–ª—è –Ω–µ–¥–æ—Å—Ç–∞—é—â–∏—Ö –¥–∞–Ω–Ω—ã—Ö
        department_name = context.get("department", "")
        subdivision_name = context.get("subdivision", "")
        organization_name = context.get("organization_name", "")

        # –û—á–∏—â–∞–µ–º –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏–µ –æ—Ç –¥–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–æ–π –∏–Ω—Ñ–æ—Ä–º–∞—Ü–∏–∏, –µ—Å–ª–∏ –æ–Ω–∞ –µ—Å—Ç—å
        if subdivision_name and "(" in subdivision_name:
            # –û—Å—Ç–∞–≤–ª—è–µ–º —Ç–æ–ª—å–∫–æ –Ω–∞–∑–≤–∞–Ω–∏–µ –¥–æ —Å–∫–æ–±–∫–∏
            subdivision_name = subdivision_name.split("(")[0].strip()

        employee_sizes = get_employee_sizes(employee, gender)
        context.update({
            "card_number": f"SIZ-{employee.id}",
            "employee_full_name": full_name,
            "last_name": last_name,
            "first_name": first_name,
            "patronymic": patronymic,
            "employee_gender": gender,
            "employee_height": employee_sizes["height"],
            "employee_clothing_size": employee_sizes["clothing_size"],
            "employee_shoe_size": employee_sizes["shoe_size"],
            "siz_issue_date": "",
            "department_name": department_name,
            "subdivision_name": subdivision_name,  # –î–æ–±–∞–≤–ª—è–µ–º –ø–æ–¥—Ä–∞–∑–¥–µ–ª–µ–Ω–∏–µ –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç
            "organization_name": organization_name,  # –î–æ–±–∞–≤–ª—è–µ–º –æ—Ä–≥–∞–Ω–∏–∑–∞—Ü–∏—é –≤ –∫–æ–Ω—Ç–µ–∫—Å—Ç
            "position_name": context.get("position_nominative", ""),
            "hire_date": employee.hire_date.strftime("%d.%m.%Y") if hasattr(employee,
                                                                            "hire_date") and employee.hire_date else "",
            # –†–∞–∑–º–µ—Ä—ã –°–ò–ó
            "ppe_head": ppe_head,
            "ppe_gloves": ppe_gloves,
            "sizod": sizod,

            # –ú–∞—Ä–∫–µ—Ä—ã –¥–ª—è —Ç–∞–±–ª–∏—Ü
            "NORMS_TABLE": "NORMS_TABLE_MARKER",
            "ISSUED_TABLE": "ISSUED_TABLE_MARKER",

            # –î–∞–Ω–Ω—ã–µ –¥–ª—è —Ç–∞–±–ª–∏—Ü - —Ä–∞–∑–¥–µ–ª—è–µ–º –¥–ª—è –ª–∏—Ü–µ–≤–æ–π –∏ –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã
            "siz_norms": all_norms_data,  # –í–°–ï –°–ò–ó –¥–ª—è –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω—ã
            "issued_siz": selected_norms_data  # –í–´–ë–†–ê–ù–ù–´–ï –°–ò–ó –¥–ª—è –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã
        })

        # 9. –î–æ–±–∞–≤–ª–µ–Ω–∏–µ –ø–æ–ª—å–∑–æ–≤–∞—Ç–µ–ª—å—Å–∫–æ–≥–æ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞
        if custom_context:
            for k, v in custom_context.items():
                if k not in ['selected_norm_ids', 'selected_norms']:
                    context[k] = v

        # 10. –ì–µ–Ω–µ—Ä–∞—Ü–∏—è –¥–æ–∫—É–º–µ–Ω—Ç–∞ + –ø–æ—Å—Ç-–æ–±—Ä–∞–±–æ—Ç–∫–∞
        return generate_docx_from_template(
            template,
            context,
            employee,
            user,
            post_processor=process_siz_card_tables,
            raise_on_error=raise_on_error,
        )

    except Exception as exc:
        logger.error("–û—à–∏–±–∫–∞ –ø—Ä–∏ –≥–µ–Ω–µ—Ä–∞—Ü–∏–∏ –∫–∞—Ä—Ç–æ—á–∫–∏ –°–ò–ó: %s", exc)
        logger.error(traceback.format_exc())
        if raise_on_error:
            raise
        return None


# =========================
#   –ì–ï–ù–ï–†–ê–¶–ò–Ø –†–ê–ó–ú–ï–†–û–í –°–ò–ó
# =========================

def _split_full_name(full_name: str) -> Tuple[str, str, str]:
    """–†–∞–∑–¥–µ–ª—è–µ—Ç –ø–æ–ª–Ω–æ–µ –∏–º—è –Ω–∞ —Ñ–∞–º–∏–ª–∏—é, –∏–º—è –∏ –æ—Ç—á–µ—Å—Ç–≤–æ."""
    parts = full_name.split()

    if len(parts) >= 3:
        return parts[0], parts[1], parts[2]
    elif len(parts) == 2:
        return parts[0], parts[1], ""
    elif len(parts) == 1:
        return parts[0], "", ""
    else:
        return "", "", ""


def _generate_random_ppe_sizes(gender: str) -> Tuple[str, str, str, str]:
    """
    –ì–µ–Ω–µ—Ä–∏—Ä—É–µ—Ç —Å–ª—É—á–∞–π–Ω—ã–µ —Ä–∞–∑–º–µ—Ä—ã –°–ò–ó –≤ –∑–∞–≤–∏—Å–∏–º–æ—Å—Ç–∏ –æ—Ç –ø–æ–ª–∞.

    Args:
        gender: –ü–æ–ª —Å–æ—Ç—Ä—É–¥–Ω–∏–∫–∞ ("–ú—É–∂—Å–∫–æ–π" –∏–ª–∏ "–ñ–µ–Ω—Å–∫–∏–π")

    Returns:
        –ö–æ—Ä—Ç–µ–∂ (headgear, gloves, respirator, gas_mask)
    """
    if gender == "–ú—É–∂—Å–∫–æ–π":
        # –ú—É–∂—Å–∫–∏–µ —Ä–∞–∑–º–µ—Ä—ã
        headgear = random.randint(55, 59)  # –ì–æ–ª–æ–≤–Ω–æ–π —É–±–æ—Ä –æ—Ç 55 –¥–æ 59
        gloves = random.randint(15, 19) / 2  # –ü–µ—Ä—á–∞—Ç–∫–∏ –æ—Ç 7.5 –¥–æ 9.5, –∫—Ä–∞—Ç–Ω—ã–µ 0.5
        respirator = random.choice(["1", "2", "3"])  # –†–µ—Å–ø–∏—Ä–∞—Ç–æ—Ä —Ä–∞–∑–º–µ—Ä—ã 1, 2, 3
    else:
        # –ñ–µ–Ω—Å–∫–∏–µ —Ä–∞–∑–º–µ—Ä—ã
        headgear = random.randint(53, 57)  # –ì–æ–ª–æ–≤–Ω–æ–π —É–±–æ—Ä –æ—Ç 53 –¥–æ 57
        gloves = random.randint(13, 17) / 2  # –ü–µ—Ä—á–∞—Ç–∫–∏ –æ—Ç 6.5 –¥–æ 8.5, –∫—Ä–∞—Ç–Ω—ã–µ 0.5
        respirator = random.choice(["1", "2", "3"])  # –†–µ—Å–ø–∏—Ä–∞—Ç–æ—Ä —Ä–∞–∑–º–µ—Ä—ã 1, 2, 3

    # –ü—Ä–æ—Ç–∏–≤–æ–≥–∞–∑ —Ç–∞–∫–æ–≥–æ –∂–µ —Ä–∞–∑–º–µ—Ä–∞, –∫–∞–∫ –∏ —Ä–µ—Å–ø–∏—Ä–∞—Ç–æ—Ä
    gas_mask = respirator

    return str(headgear), str(gloves), respirator, gas_mask


def _gender_from_patronymic(patronymic: str) -> str:
    """–û–ø—Ä–µ–¥–µ–ª—è–µ—Ç –ø–æ–ª –ø–æ –æ—Ç—á–µ—Å—Ç–≤—É."""
    if not patronymic:
        return "–ú—É–∂—Å–∫–æ–π"  # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é

    if patronymic.endswith(("–Ω–∞", "–≤–Ω–∞", "—á–Ω–∞", "–∫—ã–∑—ã", "–∑—ã")):
        return "–ñ–µ–Ω—Å–∫–∏–π"
    if patronymic.endswith(("–∏—á", "—ã—á", "–æ–≥–ª—ã", "—É–ª—ã", "–ª—ã")):
        return "–ú—É–∂—Å–∫–æ–π"

    # –ü–æ —É–º–æ–ª—á–∞–Ω–∏—é —Å—á–∏—Ç–∞–µ–º –º—É–∂—Å–∫–∏–º
    return "–ú—É–∂—Å–∫–æ–π"


# =========================
#   –ü–û–°–¢-–û–ë–†–ê–ë–û–¢–ö–ê –¢–ê–ë–õ–ò–¶
# =========================

def process_siz_card_tables(doc, context):
    """–û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç —Ç–∞–±–ª–∏—Ü—ã –≤ —Å–≥–µ–Ω–µ—Ä–∏—Ä–æ–≤–∞–Ω–Ω–æ–º –¥–æ–∫—É–º–µ–Ω—Ç–µ."""
    try:
        logger.info("=== –ù–ê–ß–ê–õ–û –ü–û–°–¢-–û–ë–†–ê–ë–û–¢–ö–ò –¢–ê–ë–õ–ò–¶ –°–ò–ó ===")
        docx_document = doc.docx

        logger.info(f"–ö–æ–ª–∏—á–µ—Å—Ç–≤–æ —Ç–∞–±–ª–∏—Ü –≤ –¥–æ–∫—É–º–µ–Ω—Ç–µ: {len(docx_document.tables)}")
        logger.info(f"–ö–ª—é—á–∏ –∫–æ–Ω—Ç–µ–∫—Å—Ç–∞: {list(context.keys())}")

        # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –ª–∏—Ü–µ–≤—É—é —Å—Ç–æ—Ä–æ–Ω—É —Å –ü–û–õ–ù–´–ú —Å–ø–∏—Å–∫–æ–º –Ω–æ—Ä–º
        norms_data = context.get("siz_norms", [])
        logger.info(f"–ü–æ–ª—É—á–µ–Ω–æ –Ω–æ—Ä–º –¥–ª—è –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω—ã: {len(norms_data)}")

        if norms_data:
            processed_front = False
            for i, table in enumerate(docx_document.tables):
                logger.info(f"–ü—Ä–æ–≤–µ—Ä–∫–∞ —Ç–∞–±–ª–∏—Ü—ã #{i} –Ω–∞ –Ω–∞–ª–∏—á–∏–µ NORMS_TABLE_MARKER")
                row_idx, cell_idx = _find_marker_in_table(
                    table,
                    ["NORMS_TABLE_MARKER", "NORMS_TABLE"]
                )
                if row_idx is not None:
                    logger.info(f"–ú–∞—Ä–∫–µ—Ä NORMS_TABLE_MARKER –Ω–∞–π–¥–µ–Ω –≤ —Ç–∞–±–ª–∏—Ü–µ #{i}, —Å—Ç—Ä–æ–∫–∞ {row_idx}, —è—á–µ–π–∫–∞ {cell_idx}")
                    # –û–±—Ä–∞–±–æ—Ç–∫–∞ —Ç–∞–±–ª–∏—Ü—ã –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω—ã
                    process_front_table(table, row_idx, cell_idx, norms_data)
                    processed_front = True
                    break

            if not processed_front:
                logger.warning("–ú–∞—Ä–∫–µ—Ä NORMS_TABLE_MARKER –Ω–µ –Ω–∞–π–¥–µ–Ω –≤ —Ç–∞–±–ª–∏—Ü–∞—Ö")
                # –î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –¥–∏–∞–≥–Ω–æ—Å—Ç–∏–∫–∞
                for i, table in enumerate(docx_document.tables):
                    logger.warning(f"–¢–∞–±–ª–∏—Ü–∞ #{i}, –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —Å—Ç—Ä–æ–∫: {len(table.rows)}")
                    for r_idx, row in enumerate(table.rows[:3]):  # –ü–µ—Ä–≤—ã–µ 3 —Å—Ç—Ä–æ–∫–∏
                        for c_idx, cell in enumerate(row.cells):
                            logger.warning(f"  –°—Ç—Ä–æ–∫–∞ {r_idx}, —è—á–µ–π–∫–∞ {c_idx}: '{cell.text[:50]}'")

        # –û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ–º –æ–±–æ—Ä–æ—Ç–Ω—É—é —Å—Ç–æ—Ä–æ–Ω—É —Ç–æ–ª—å–∫–æ —Å –í–´–ë–†–ê–ù–ù–´–ú–ò –Ω–æ—Ä–º–∞–º–∏
        issued_data = context.get("issued_siz", [])
        issue_date = context.get("siz_issue_date", "") or ""
        logger.info(f"–ü–æ–ª—É—á–µ–Ω–æ –≤—ã–±—Ä–∞–Ω–Ω—ã—Ö –Ω–æ—Ä–º –¥–ª—è –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã: {len(issued_data)}")

        if issued_data:
            processed_back = False
            for i, table in enumerate(docx_document.tables):
                logger.info(f"–ü—Ä–æ–≤–µ—Ä–∫–∞ —Ç–∞–±–ª–∏—Ü—ã #{i} –Ω–∞ –Ω–∞–ª–∏—á–∏–µ ISSUED_TABLE_MARKER")
                row_idx, cell_idx = _find_marker_in_table(
                    table,
                    ["ISSUED_TABLE_MARKER", "ISSUED_TABLE"]
                )
                if row_idx is not None:
                    logger.info(f"–ú–∞—Ä–∫–µ—Ä ISSUED_TABLE_MARKER –Ω–∞–π–¥–µ–Ω –≤ —Ç–∞–±–ª–∏—Ü–µ #{i}, —Å—Ç—Ä–æ–∫–∞ {row_idx}, —è—á–µ–π–∫–∞ {cell_idx}")
                    # –û–±—Ä–∞–±–æ—Ç–∫–∞ —Ç–∞–±–ª–∏—Ü—ã –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã
                    process_back_table(table, row_idx, cell_idx, issued_data, issue_date)
                    processed_back = True
                    break

            if not processed_back:
                logger.warning("–ú–∞—Ä–∫–µ—Ä ISSUED_TABLE_MARKER –Ω–µ –Ω–∞–π–¥–µ–Ω –≤ —Ç–∞–±–ª–∏—Ü–∞—Ö")

        logger.info("=== –ö–û–ù–ï–¶ –ü–û–°–¢-–û–ë–†–ê–ë–û–¢–ö–ò –¢–ê–ë–õ–ò–¶ –°–ò–ó ===")
        return doc

    except Exception as exc:
        logger.error("–û—à–∏–±–∫–∞ –ø—Ä–∏ –ø–æ—Å—Ç-–æ–±—Ä–∞–±–æ—Ç–∫–µ —Ç–∞–±–ª–∏—Ü: %s", exc)
        logger.error(traceback.format_exc())
        return doc


def process_front_table(table, row_idx, cell_idx, norms_data):
    """–û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç —Ç–∞–±–ª–∏—Ü—É –Ω–∞ –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω–µ —Å –æ–±—ä–µ–¥–∏–Ω—ë–Ω–Ω—ã–º–∏ —Å—Ç—Ä–æ–∫–∞–º–∏ –¥–ª—è —É—Å–ª–æ–≤–∏–π."""
    try:
        # –ì—Ä—É–ø–ø–∏—Ä—É–µ–º –Ω–æ—Ä–º—ã –ø–æ —É—Å–ª–æ–≤–∏—è–º
        grouped_norms = {}
        for norm in norms_data:
            condition = norm.get("condition", "")
            if condition not in grouped_norms:
                grouped_norms[condition] = []
            grouped_norms[condition].append(norm)

        # –£–¥–∞–ª—è–µ–º –º–∞—Ä–∫–µ—Ä –∏–∑ —è—á–µ–π–∫–∏
        cell = table.rows[row_idx].cells[cell_idx]
        cell.text = ""

        # –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º –∑–∞–≥–æ–ª–æ–≤–∫–∏ —Ç–∞–±–ª–∏—Ü—ã - —É–º–µ–Ω—å—à–∞–µ–º —Ä–∞–∑–º–µ—Ä —à—Ä–∏—Ñ—Ç–∞ –¥–æ 9 –ø—Ç
        if row_idx > 0:  # –ü—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ–º, —á—Ç–æ –∑–∞–≥–æ–ª–æ–≤–æ–∫ –Ω–∞—Ö–æ–¥–∏—Ç—Å—è –≤ —Å—Ç—Ä–æ–∫–µ –≤—ã—à–µ –º–∞—Ä–∫–µ—Ä–∞
            header_row = table.rows[row_idx - 1]
            _format_header_row(header_row)

        # –ó–∞–ø–æ–ª–Ω—è–µ–º –ø–µ—Ä–≤—É—é —Å—Ç—Ä–æ–∫—É, –µ—Å–ª–∏ –µ—Å—Ç—å –Ω–æ—Ä–º—ã –±–µ–∑ —É—Å–ª–æ–≤–∏–π
        current_row = row_idx
        template_row = None
        if current_row < len(table.rows):
            template_row = table.rows[current_row]

        if "" in grouped_norms and grouped_norms[""]:
            for norm in grouped_norms[""]:
                if current_row >= len(table.rows):
                    new_row = table.add_row()
                    # –ö–æ–ø–∏—Ä—É–µ–º —Ñ–æ—Ä–º–∞—Ç —Å –æ–±—Ä–∞–∑—Ü–∞, –µ—Å–ª–∏ –æ–Ω –µ—Å—Ç—å
                    if template_row:
                        _copy_row_properties(template_row, new_row)
                else:
                    new_row = table.rows[current_row]

                _fill_front_row(new_row, norm)
                current_row += 1

        # –î–æ–±–∞–≤–ª—è–µ–º —Å—Ç—Ä–æ–∫–∏ —Å —É—Å–ª–æ–≤–∏—è–º–∏ –∏ —Å–æ–æ—Ç–≤–µ—Ç—Å—Ç–≤—É—é—â–∏–º–∏ –Ω–æ—Ä–º–∞–º–∏
        for condition, norms in grouped_norms.items():
            if not condition:
                continue  # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –ø—É—Å—Ç—ã–µ —É—Å–ª–æ–≤–∏—è, –æ–Ω–∏ —É–∂–µ –æ–±—Ä–∞–±–æ—Ç–∞–Ω—ã

            # –î–æ–±–∞–≤–ª—è–µ–º —Å—Ç—Ä–æ–∫—É —Å —É—Å–ª–æ–≤–∏–µ–º
            if current_row >= len(table.rows):
                condition_row = table.add_row()
                # –ï—Å–ª–∏ –µ—Å—Ç—å —à–∞–±–ª–æ–Ω–Ω–∞—è —Å—Ç—Ä–æ–∫–∞, –∫–æ–ø–∏—Ä—É–µ–º –µ—ë —Å–≤–æ–π—Å—Ç–≤–∞
                if template_row:
                    _copy_row_properties(template_row, condition_row)
            else:
                condition_row = table.rows[current_row]

            # –û–±—ä–µ–¥–∏–Ω—è–µ–º —è—á–µ–π–∫–∏ –≤ —Å—Ç—Ä–æ–∫–µ —É—Å–ª–æ–≤–∏—è
            first_cell = condition_row.cells[0]
            for i in range(1, len(condition_row.cells)):
                if i < len(condition_row.cells):
                    try:
                        first_cell.merge(condition_row.cells[i])
                    except Exception as e:
                        logger.warning(f"–ù–µ —É–¥–∞–ª–æ—Å—å –æ–±—ä–µ–¥–∏–Ω–∏—Ç—å —è—á–µ–π–∫–∏: {str(e)}")

            # –ó–∞–ø–æ–ª–Ω—è–µ–º —Ç–µ–∫—Å—Ç —É—Å–ª–æ–≤–∏—è
            first_cell.text = condition
            for paragraph in first_cell.paragraphs:
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–µ–∂–¥—É—Å—Ç—Ä–æ—á–Ω—ã–π –∏–Ω—Ç–µ—Ä–≤–∞–ª –∫–∞–∫ –æ–¥–∏–Ω–∞—Ä–Ω—ã–π
                paragraph.paragraph_format.line_spacing = 1.0
                # –£–±–∏—Ä–∞–µ–º –æ—Ç—Å—Ç—É–ø—ã –¥–æ –∏ –ø–æ—Å–ª–µ –∞–±–∑–∞—Ü–∞
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)
                    run.italic = True

            current_row += 1

            # –î–æ–±–∞–≤–ª—è–µ–º –Ω–æ—Ä–º—ã –¥–ª—è –¥–∞–Ω–Ω–æ–≥–æ —É—Å–ª–æ–≤–∏—è
            for norm in norms:
                if current_row >= len(table.rows):
                    new_row = table.add_row()
                    # –ö–æ–ø–∏—Ä—É–µ–º —Ñ–æ—Ä–º–∞—Ç —Å –æ–±—Ä–∞–∑—Ü–∞, –µ—Å–ª–∏ –æ–Ω –µ—Å—Ç—å
                    if template_row:
                        _copy_row_properties(template_row, new_row)
                else:
                    new_row = table.rows[current_row]

                _fill_front_row(new_row, norm)
                current_row += 1

        # –ü—Ä–∏–º–µ–Ω—è–µ–º —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∫–æ –≤—Å–µ–π —Ç–∞–±–ª–∏—Ü–µ
        _apply_table_format(table)

    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ —Ç–∞–±–ª–∏—Ü—ã –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω—ã: {str(e)}")
        logger.error(traceback.format_exc())


def process_back_table(table, row_idx, cell_idx, norms_data, issue_date=""):
    """–û–±—Ä–∞–±–∞—Ç—ã–≤–∞–µ—Ç —Ç–∞–±–ª–∏—Ü—É –Ω–∞ –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω–µ, –∑–∞–ø–æ–ª–Ω—è—è —Ç–æ–ª—å–∫–æ –Ω—É–∂–Ω—ã–µ –∫–æ–ª–æ–Ω–∫–∏."""
    try:
        # –£–¥–∞–ª—è–µ–º –º–∞—Ä–∫–µ—Ä –∏–∑ —è—á–µ–π–∫–∏
        cell = table.rows[row_idx].cells[cell_idx]
        cell.text = ""

        # –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º –∑–∞–≥–æ–ª–æ–≤–∫–∏ —Ç–∞–±–ª–∏—Ü—ã - —É–º–µ–Ω—å—à–∞–µ–º —Ä–∞–∑–º–µ—Ä —à—Ä–∏—Ñ—Ç–∞ –¥–æ 9 –ø—Ç
        if row_idx > 0:  # –ü—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ–º, —á—Ç–æ –∑–∞–≥–æ–ª–æ–≤–æ–∫ –Ω–∞—Ö–æ–¥–∏—Ç—Å—è –≤ —Å—Ç—Ä–æ–∫–µ –≤—ã—à–µ –º–∞—Ä–∫–µ—Ä–∞
            header_row = table.rows[row_idx - 1]
            _format_header_row(header_row)

            # –ï—Å–ª–∏ –µ—Å—Ç—å —Å—Ç—Ä–æ–∫–∞ —Å –ø–æ–¥–∑–∞–≥–æ–ª–æ–≤–∫–∞–º–∏, —Ñ–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º –∏ –µ—ë
            if row_idx > 1:
                subheader_row = table.rows[row_idx - 2]
                _format_header_row(subheader_row)

        # –û–ø—Ä–µ–¥–µ–ª—è–µ–º –∫–æ–ª–æ–Ω–∫–∏ –¥–ª—è –∑–∞–ø–æ–ª–Ω–µ–Ω–∏—è (0-based): 0, 1, 2, 3, 6
        cols_to_fill = [0, 1, 2, 3, 5, 6]

        # –°–æ–∑–¥–∞–µ–º –æ–±—Ä–∞–∑–µ—Ü —Å—Ç—Ä–æ–∫–∏ —Å –ø—Ä–∞–≤–∏–ª—å–Ω—ã–º —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ–º
        template_row = None
        if row_idx < len(table.rows):
            template_row = table.rows[row_idx]

        # –ó–∞–ø–æ–ª–Ω—è–µ–º —Å—Ç—Ä–æ–∫–∏ –¥–∞–Ω–Ω—ã–º–∏
        current_row = row_idx
        for norm in norms_data:
            if current_row >= len(table.rows):
                new_row = table.add_row()
                # –ö–æ–ø–∏—Ä—É–µ–º —Ñ–æ—Ä–º–∞—Ç —Å –æ–±—Ä–∞–∑—Ü–∞, –µ—Å–ª–∏ –æ–Ω –µ—Å—Ç—å
                if template_row:
                    _copy_row_properties(template_row, new_row)
            else:
                new_row = table.rows[current_row]

            # –ó–∞–ø–æ–ª–Ω—è–µ–º —Ç–æ–ª—å–∫–æ –Ω—É–∂–Ω—ã–µ –∫–æ–ª–æ–Ω–∫–∏
            if 0 < len(new_row.cells):
                new_row.cells[0].text = norm.get("name", "")
                # –í—ã—Ä–∞–≤–Ω–∏–≤–∞–Ω–∏–µ –ø–æ –ª–µ–≤–æ–º—É –∫—Ä–∞—é –¥–ª—è –ø–µ—Ä–≤–æ–π —è—á–µ–π–∫–∏
                for paragraph in new_row.cells[0].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 1 < len(new_row.cells):
                new_row.cells[1].text = norm.get("classification", "")
                # –¶–µ–Ω—Ç—Ä–∏—Ä–æ–≤–∞–Ω–∏–µ –¥–ª—è –æ—Å—Ç–∞–ª—å–Ω—ã—Ö —è—á–µ–µ–∫
                for paragraph in new_row.cells[1].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 2 < len(new_row.cells):
                new_row.cells[2].text = issue_date
                for paragraph in new_row.cells[2].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 3 < len(new_row.cells):
                new_row.cells[3].text = str(norm.get("quantity", ""))
                for paragraph in new_row.cells[3].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 5 < len(new_row.cells):
                new_row.cells[5].text = norm.get("cost", "")
                for paragraph in new_row.cells[5].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            if 6 < len(new_row.cells):
                new_row.cells[6].text = "‚úì"  # –ì–∞–ª–æ—á–∫–∞ –≤ –∫–æ–ª–æ–Ω–∫–µ —Ä–∞—Å–ø–∏—Å–∫–∏
                for paragraph in new_row.cells[6].paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.0
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = "Times New Roman"
                        run.font.size = Pt(12)

            # –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º —è—á–µ–π–∫–∏
            for col in cols_to_fill:
                if col < len(new_row.cells):
                    _set_cell_border(new_row.cells[col])

            current_row += 1

        # –ü—Ä–∏–º–µ–Ω—è–µ–º —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∫–æ –≤—Å–µ–π —Ç–∞–±–ª–∏—Ü–µ
        _apply_table_format(table)

    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –æ–±—Ä–∞–±–æ—Ç–∫–µ —Ç–∞–±–ª–∏—Ü—ã –æ–±–æ—Ä–æ—Ç–Ω–æ–π —Å—Ç–æ—Ä–æ–Ω—ã: {str(e)}")
        logger.error(traceback.format_exc())


# --------------------------
#   –£–¢–ò–õ–ò–¢–´
# --------------------------

def _copy_row_properties(src_row, dst_row):
    """–ö–æ–ø–∏—Ä—É–µ—Ç –≤—Å–µ —Å–≤–æ–π—Å—Ç–≤–∞ –∏–∑ –∏—Å—Ö–æ–¥–Ω–æ–π —Å—Ç—Ä–æ–∫–∏ –≤ —Ü–µ–ª–µ–≤—É—é."""
    try:
        # –ö–æ–ø–∏—Ä—É–µ–º –≤—ã—Å–æ—Ç—É —Å—Ç—Ä–æ–∫–∏ –∏ –¥—Ä—É–≥–∏–µ –∞—Ç—Ä–∏–±—É—Ç—ã
        dst_row.height = src_row.height

        # –ö–æ–ø–∏—Ä—É–µ–º –∫–æ–ª–∏—á–µ—Å—Ç–≤–æ —è—á–µ–µ–∫ –∏ –∏—Ö —Å–≤–æ–π—Å—Ç–≤–∞
        for i in range(min(len(src_row.cells), len(dst_row.cells))):
            # –°–±—Ä–∞—Å—ã–≤–∞–µ–º —Ç–µ–∫—Å—Ç, –Ω–æ –∫–æ–ø–∏—Ä—É–µ–º —Ñ–æ—Ä–º–∞—Ç
            dst_row.cells[i].text = ""

            # –ü—ã—Ç–∞–µ–º—Å—è –∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å—Ç–∏–ª—å —è—á–µ–π–∫–∏ –∏ –µ—ë –≥—Ä–∞–Ω–∏—Ü—ã
            try:
                dst_tc = dst_row.cells[i]._tc
                src_tc = src_row.cells[i]._tc

                # –ö–æ–ø–∏—Ä—É–µ–º –≤—Å–µ —Å–≤–æ–π—Å—Ç–≤–∞ —Ç–∞–±–ª–∏—Ü—ã
                if src_tc.tcPr is not None:
                    if dst_tc.tcPr is None:
                        dst_tc.tcPr = OxmlElement('w:tcPr')

                    # –ö–æ–ø–∏—Ä—É–µ–º —à–∏—Ä–∏–Ω—É —è—á–µ–π–∫–∏
                    src_width = src_tc.tcPr.find(qn('w:tcW'))
                    if src_width is not None:
                        dst_width = dst_tc.tcPr.find(qn('w:tcW'))
                        if dst_width is None:
                            dst_width = OxmlElement('w:tcW')
                            dst_tc.tcPr.append(dst_width)
                        dst_width.attrib.update(src_width.attrib)
            except Exception as e:
                logger.debug(f"–ù–µ —É–¥–∞–ª–æ—Å—å —Å–∫–æ–ø–∏—Ä–æ–≤–∞—Ç—å —Å–≤–æ–π—Å—Ç–≤–∞ —è—á–µ–π–∫–∏: {str(e)}")
                # –ü—Ä–æ–¥–æ–ª–∂–∞–µ–º –≤—ã–ø–æ–ª–Ω–µ–Ω–∏–µ

            # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –≥—Ä–∞–Ω–∏—Ü—ã
            _set_cell_border(dst_row.cells[i])
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ –∫–æ–ø–∏—Ä–æ–≤–∞–Ω–∏–∏ —Å–≤–æ–π—Å—Ç–≤ —Å—Ç—Ä–æ–∫–∏: {str(e)}")


def _format_header_row(header_row):
    """–§–æ—Ä–º–∞—Ç–∏—Ä—É–µ—Ç —Å—Ç—Ä–æ–∫—É –∑–∞–≥–æ–ª–æ–≤–∫–∞ —Ç–∞–±–ª–∏—Ü—ã."""
    try:
        for cell in header_row.cells:
            for paragraph in cell.paragraphs:
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–µ–∂–¥—É—Å—Ç—Ä–æ—á–Ω—ã–π –∏–Ω—Ç–µ—Ä–≤–∞–ª –∫–∞–∫ –æ–¥–∏–Ω–∞—Ä–Ω—ã–π
                paragraph.paragraph_format.line_spacing = 1.0
                # –£–±–∏—Ä–∞–µ–º –æ—Ç—Å—Ç—É–ø—ã –¥–æ –∏ –ø–æ—Å–ª–µ –∞–±–∑–∞—Ü–∞
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)

                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(9)  # –£–º–µ–Ω—å—à–∞–µ–º —Ä–∞–∑–º–µ—Ä –¥–æ 9 –ø—É–Ω–∫—Ç–æ–≤
                    run.font.bold = True  # –î–µ–ª–∞–µ–º –∑–∞–≥–æ–ª–æ–≤–∫–∏ –ø–æ–ª—É–∂–∏—Ä–Ω—ã–º–∏
    except Exception as e:
        logger.error(f"–û—à–∏–±–∫–∞ –ø—Ä–∏ —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–∏ –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤: {str(e)}")


def _find_marker_in_table(table, markers):
    """–ù–∞—Ö–æ–¥–∏—Ç –º–∞—Ä–∫–µ—Ä –≤ —Ç–∞–±–ª–∏—Ü–µ –∏ –≤–æ–∑–≤—Ä–∞—â–∞–µ—Ç –∫–æ–æ—Ä–¥–∏–Ω–∞—Ç—ã —è—á–µ–π–∫–∏.

    –ü—Ä–æ–≤–µ—Ä—è–µ—Ç –∫–∞–∫ –ø–æ–ª–Ω—ã–π —Ç–µ–∫—Å—Ç –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞ (paragraph.text), —Ç–∞–∫ –∏ —Ç–µ–∫—Å—Ç –∫–∞–∂–¥–æ–≥–æ run,
    —á—Ç–æ–±—ã –Ω–∞–π—Ç–∏ –º–∞—Ä–∫–µ—Ä –¥–∞–∂–µ –µ—Å–ª–∏ –æ–Ω —Ä–∞–∑–±–∏—Ç –Ω–∞ –Ω–µ—Å–∫–æ–ª—å–∫–æ run'–æ–≤ –≤ XML.
    """
    if isinstance(markers, str):
        markers = [markers]

    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            # –ü—Ä–æ–≤–µ—Ä—è–µ–º –≤–µ—Å—å —Ç–µ–∫—Å—Ç —è—á–µ–π–∫–∏
            cell_full_text = cell.text
            for marker in markers:
                if marker in cell_full_text:
                    logger.debug(f"–ú–∞—Ä–∫–µ—Ä '{marker}' –Ω–∞–π–¥–µ–Ω –≤ cell.text: '{cell_full_text[:100]}'")
                    return r_idx, c_idx

            # –î–æ–ø–æ–ª–Ω–∏—Ç–µ–ª—å–Ω–∞—è –ø—Ä–æ–≤–µ—Ä–∫–∞ –ø–æ –ø–∞—Ä–∞–≥—Ä–∞—Ñ–∞–º
            for paragraph in cell.paragraphs:
                for marker in markers:
                    if marker in paragraph.text:
                        logger.debug(f"–ú–∞—Ä–∫–µ—Ä '{marker}' –Ω–∞–π–¥–µ–Ω –≤ paragraph.text: '{paragraph.text[:100]}'")
                        return r_idx, c_idx

                # –ü—Ä–æ–≤–µ—Ä—è–µ–º –∫–∞–∂–¥—ã–π run –æ—Ç–¥–µ–ª—å–Ω–æ (–Ω–∞ —Å–ª—É—á–∞–π –µ—Å–ª–∏ –º–∞—Ä–∫–µ—Ä —Ä–∞–∑–±–∏—Ç)
                full_run_text = ''.join(run.text for run in paragraph.runs)
                for marker in markers:
                    if marker in full_run_text:
                        logger.debug(f"–ú–∞—Ä–∫–µ—Ä '{marker}' –Ω–∞–π–¥–µ–Ω –≤ –æ–±—ä–µ–¥–∏–Ω—ë–Ω–Ω—ã—Ö runs: '{full_run_text[:100]}'")
                        return r_idx, c_idx

    logger.debug(f"–ú–∞—Ä–∫–µ—Ä—ã '{markers}' –Ω–µ –Ω–∞–π–¥–µ–Ω—ã –Ω–∏ –≤ –æ–¥–Ω–æ–π —è—á–µ–π–∫–µ —Ç–∞–±–ª–∏—Ü—ã")
    return None, None


def _fill_front_row(row, norm):
    """–ó–∞–ø–æ–ª–Ω—è–µ—Ç —Å—Ç—Ä–æ–∫—É —Ç–∞–±–ª–∏—Ü—ã –Ω–∞ –ª–∏—Ü–µ–≤–æ–π —Å—Ç–æ—Ä–æ–Ω–µ."""
    # –ü—Ä–æ–≤–µ—Ä—è–µ–º –Ω–∞–ª–∏—á–∏–µ –¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ–≥–æ –∫–æ–ª–∏—á–µ—Å—Ç–≤–∞ —è—á–µ–µ–∫
    if len(row.cells) < 5:
        logger.warning(f"–ù–µ–¥–æ—Å—Ç–∞—Ç–æ—á–Ω–æ —è—á–µ–µ–∫ –≤ —Å—Ç—Ä–æ–∫–µ —Ç–∞–±–ª–∏—Ü—ã: {len(row.cells)}")
        return

    # –û—á–∏—â–∞–µ–º —è—á–µ–π–∫–∏ –ø–µ—Ä–µ–¥ –∑–∞–ø–æ–ª–Ω–µ–Ω–∏–µ–º –¥–ª—è –ø—Ä–µ–¥–æ—Ç–≤—Ä–∞—â–µ–Ω–∏—è –ø—Ä–æ–±–ª–µ–º —Å —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ–º
    for i in range(5):
        row.cells[i].text = ""

    # –ó–∞–ø–æ–ª–Ω—è–µ–º —è—á–µ–π–∫–∏
    row.cells[0].text = norm.get("name", "")
    row.cells[1].text = norm.get("classification", "")
    row.cells[2].text = norm.get("unit", "")
    row.cells[3].text = str(norm.get("quantity", ""))
    row.cells[4].text = norm.get("wear_period", "")

    # –§–æ—Ä–º–∞—Ç–∏—Ä—É–µ–º —è—á–µ–π–∫–∏
    for i in range(5):
        for paragraph in row.cells[i].paragraphs:
            # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–µ–∂–¥—É—Å—Ç—Ä–æ—á–Ω—ã–π –∏–Ω—Ç–µ—Ä–≤–∞–ª –∫–∞–∫ –æ–¥–∏–Ω–∞—Ä–Ω—ã–π
            paragraph.paragraph_format.line_spacing = 1.0
            # –£–±–∏—Ä–∞–µ–º –æ—Ç—Å—Ç—É–ø—ã –¥–æ –∏ –ø–æ—Å–ª–µ –∞–±–∑–∞—Ü–∞
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)

            # –ü–µ—Ä–≤—ã–π —Å—Ç–æ–ª–±–µ—Ü —Å –Ω–∞–∏–º–µ–Ω–æ–≤–∞–Ω–∏–µ–º –°–ò–ó –≤—ã—Ä–∞–≤–Ω–∏–≤–∞–µ–º –ø–æ –ª–µ–≤–æ–º—É –∫—Ä–∞—é
            if i == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                # –û—Å—Ç–∞–ª—å–Ω—ã–µ —Å—Ç–æ–ª–±—Ü—ã - –ø–æ —Ü–µ–Ω—Ç—Ä—É
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º —à—Ä–∏—Ñ—Ç –¥–ª—è –≤—Å–µ–≥–æ —Å–æ–¥–µ—Ä–∂–∏–º–æ–≥–æ —è—á–µ–π–∫–∏
            for run in paragraph.runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)

        # –ü—Ä–∏–º–µ–Ω—è–µ–º –≥—Ä–∞–Ω–∏—Ü—ã
        _set_cell_border(row.cells[i])


def _apply_table_format(table):
    """–ü—Ä–∏–º–µ–Ω—è–µ—Ç —Ñ–æ—Ä–º–∞—Ç–∏—Ä–æ–≤–∞–Ω–∏–µ –∫ —Ç–∞–±–ª–∏—Ü–µ."""
    # 1. –ü–æ–ø—ã—Ç–∫–∞ –ø—Ä–∏–º–µ–Ω–∏—Ç—å –≥–æ—Ç–æ–≤—ã–π —Å—Ç–∏–ª—å
    try:
        table.style = "Table Grid"
    except (KeyError, ValueError):
        logger.info("–°—Ç–∏–ª—å 'Table Grid' –Ω–µ–¥–æ—Å—Ç—É–ø–µ–Ω ‚Äì –ø—Ä–æ—Å—Ç–∞–≤–ª—è–µ–º –≥—Ä–∞–Ω–∏—Ü—ã –≤—Ä—É—á–Ω—É—é")
        _add_borders_manually(table)

    # 2. –®—Ä–∏—Ñ—Ç –≤—Å–µ–º run‚Äë–∞–º, –∫—Ä–æ–º–µ –∑–∞–≥–æ–ª–æ–≤–∫–æ–≤, –∫–æ—Ç–æ—Ä—ã–µ —Ñ–æ—Ä–º–∞—Ç–∏—Ä—É—é—Ç—Å—è –æ—Ç–¥–µ–ª—å–Ω–æ
    for row_idx, row in enumerate(table.rows):
        # –ü—Ä–æ–ø—É—Å–∫–∞–µ–º –∑–∞–≥–æ–ª–æ–≤–∫–∏, –∫–æ—Ç–æ—Ä—ã–µ —Ñ–æ—Ä–º–∞—Ç–∏—Ä—É—é—Ç—Å—è –æ—Ç–¥–µ–ª—å–Ω–æ
        if row_idx < 2:  # –ü—Ä–µ–¥–ø–æ–ª–∞–≥–∞–µ–º, —á—Ç–æ –∑–∞–≥–æ–ª–æ–≤–∫–∏ –≤ –ø–µ—Ä–≤—ã—Ö –¥–≤—É—Ö —Å—Ç—Ä–æ–∫–∞—Ö
            continue

        for cell in row.cells:
            for para in cell.paragraphs:
                # –£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ–º –º–µ–∂–¥—É—Å—Ç—Ä–æ—á–Ω—ã–π –∏–Ω—Ç–µ—Ä–≤–∞–ª –∫–∞–∫ –æ–¥–∏–Ω–∞—Ä–Ω—ã–π
                para.paragraph_format.line_spacing = 1.0
                # –£–±–∏—Ä–∞–µ–º –æ—Ç—Å—Ç—É–ø—ã –¥–æ –∏ –ø–æ—Å–ª–µ –∞–±–∑–∞—Ü–∞
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(0)

                if not para.runs:
                    para.add_run("")
                for run in para.runs:
                    run.font.name = "Times New Roman"
                    run.font.size = Pt(12)


def _add_borders_manually(table):
    """–î–æ–±–∞–≤–ª—è–µ—Ç –≥—Ä–∞–Ω–∏—Ü—ã –≤—Å–µ–º —è—á–µ–π–∫–∞–º —Ç–∞–±–ª–∏—Ü—ã."""
    for row in table.rows:
        for cell in row.cells:
            _set_cell_border(cell)


def _set_cell_border(
        cell,
        **kwargs,
):
    """–£—Å—Ç–∞–Ω–∞–≤–ª–∏–≤–∞–µ—Ç –æ–¥–∏–Ω–∞—Ä–Ω—É—é –≥—Ä–∞–Ω–∏—Ü—É —Ç–æ–ª—â–∏–Ω–æ–π 4 Œµ–º (‚âà0.5 pt) –≤–æ–∫—Ä—É–≥ —è—á–µ–π–∫–∏."""
    # –ó–Ω–∞—á–µ–Ω–∏—è –ø–æ —É–º–æ–ª—á–∞–Ω–∏—é ‚Äì –≤—Å–µ —Å—Ç–æ—Ä–æ–Ω—ã single 4 Œµ–º
    sides = {
        "top": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "left": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "bottom": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
        "right": {
            "val": "single",
            "sz": "4",
            "color": "000000",
            "space": "0",
        },
    }
    sides.update(kwargs)

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tblBorders = tcPr.find(qn("w:tcBorders"))
    if tblBorders is None:
        tblBorders = OxmlElement("w:tcBorders")
        tcPr.append(tblBorders)

    for edge in ("left", "top", "right", "bottom"):
        edge_el = tblBorders.find(qn(f"w:{edge}"))
        if edge_el is None:
            edge_el = OxmlElement(f"w:{edge}")
            tblBorders.append(edge_el)
        attrs = sides.get(edge, {})
        for key in ["val", "sz", "color", "space"]:
            edge_el.set(qn(f"w:{key}"), attrs.get(key, "single" if key == "val" else "4"))
