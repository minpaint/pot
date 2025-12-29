# directory/management/commands/init_document_templates.py
"""
📂 Команда для инициализации шаблонов документов

Эта команда создает начальные записи шаблонов документов в базе данных,
что позволяет использовать функциональность генерации документов
сразу после установки приложения.
"""
import os
import logging
from django.core.management.base import BaseCommand, CommandError
from django.conf import settings
from django.core.files.base import ContentFile
from django.core.files import File

from directory.models.document_template import DocumentTemplate

logger = logging.getLogger(__name__)


class Command(BaseCommand):
    help = 'Инициализирует шаблоны документов для генерации'

    def handle(self, *args, **options):
        # Создаем директорию для шаблонов, если она не существует
        templates_dir = os.path.join(settings.MEDIA_ROOT, 'document_templates')
        if not os.path.exists(templates_dir):
            os.makedirs(templates_dir)
            self.stdout.write(self.style.WARNING(f'Создана директория {templates_dir}'))

        # Создаем шаблоны, если они отсутствуют
        self._create_template(
            'all_orders',
            'Распоряжения о стажировке',
            'Шаблон распоряжений о стажировке и допуске к работе',
            'all_order_template.docx'
        )

        self._create_template(
            'knowledge_protocol',
            'Протокол проверки знаний',
            'Шаблон протокола проверки знаний по охране труда',
            'knowledge_protocol_template.docx'
        )

        self._create_template(
            'periodic_protocol',
            'Протокол периодической проверки знаний',
            'Шаблон протокола периодической проверки знаний',
            'periodic_protocol_template.docx'
        )

        self._create_template(
            'doc_familiarization',
            'Лист ознакомления с документами',
            'Шаблон листа ознакомления с документами',
            'doc_familiarization_template.docx'
        )

        self._create_template(
            'siz_card',
            'Карточка учета СИЗ',
            'Шаблон карточки учета средств индивидуальной защиты',
            'siz_card_template.docx'
        )

        self._create_template(
            'personal_ot_card',
            'Личная карточка по ОТ',
            'Шаблон личной карточки по охране труда',
            'personal_ot_card_template.docx'
        )

        self._create_template(
            'journal_example',
            'Образец заполнения журнала',
            'Шаблон образца заполнения журнала',
            'journal_example_template.docx'
        )

        self.stdout.write(self.style.SUCCESS('Шаблоны документов успешно инициализированы'))

    def _create_template(self, doc_type, name, description, filename):
        """Создает шаблон документа если он не существует"""
        # Проверяем наличие шаблона
        if DocumentTemplate.objects.filter(document_type=doc_type, is_default=True).exists():
            self.stdout.write(f'Шаблон для {doc_type} уже существует')
            return

        # Создаем объект шаблона
        template = DocumentTemplate(
            name=name,
            description=description,
            document_type=doc_type,
            is_default=True,
            is_active=True
        )

        # Проверяем, существует ли шаблон в файловой системе
        template_path = os.path.join(settings.MEDIA_ROOT, 'document_templates', filename)

        if not os.path.exists(template_path):
            # Создаем простой DOCX файл
            from docx import Document
            doc = Document()

            # Заголовок
            doc.add_heading(f"Шаблон {name}", 0)

            # Простой пример переменных
            doc.add_paragraph(f"Это шаблон для {name}.")
            doc.add_paragraph("Заменить этот файл на реальный шаблон с переменными.")

            # Сохраняем документ
            doc.save(template_path)
            self.stdout.write(f'Создан пустой шаблон DOCX для {doc_type}')

        # Сохраняем объект с привязкой к файлу
        with open(template_path, 'rb') as f:
            template.template_file.save(filename, File(f))

        template.save()
        self.stdout.write(f'Создан шаблон для {doc_type}')
