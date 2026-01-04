# directory/utils/docx_cleaner.py

import logging
from docx import Document
from io import BytesIO

logger = logging.getLogger(__name__)


def _remove_marker_from_paragraph(paragraph, marker='{# keep_empty #}'):
    """
    Удаляет маркер из параграфа, сохраняя форматирование.

    Args:
        paragraph: Параграф из которого нужно удалить маркер
        marker: Маркер для удаления
    """
    for run in paragraph.runs:
        if marker in run.text:
            run.text = run.text.replace(marker, '')


def remove_empty_paragraphs(doc_bytes: bytes) -> bytes:
    """
    Удаляет пустые параграфы из DOCX документа (body, headers, footers).

    Параграф считается пустым, если:
    - Не содержит текста
    - Или содержит только пробелы, тире и дефисы

    Args:
        doc_bytes: DOCX документ в виде байтов

    Returns:
        Очищенный DOCX документ в виде байтов
    """
    try:
        doc = Document(BytesIO(doc_bytes))

        # Удаляем пустые параграфы из body
        total_paragraphs = len(doc.paragraphs)
        paragraphs_to_remove = []

        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()

            # Проверяем наличие маркера {# keep_empty #}
            if '{# keep_empty #}' in text:
                # Удаляем маркер из текста, сохраняя форматирование
                _remove_marker_from_paragraph(paragraph)
                logger.debug(f"[remove_empty_paragraphs] Body параграф {i} сохранён с маркером keep_empty")
                continue

            # Проверяем, является ли параграф пустым
            if not text:
                paragraphs_to_remove.append(i)
                continue

            # Проверяем, состоит ли строка только из разделителей (тире и пробелы)
            # Разрешённые символы: пробел, дефис, короткое тире, длинное тире
            if all(char in ' -–—' for char in text):
                paragraphs_to_remove.append(i)
                logger.debug(f"[remove_empty_paragraphs] Body параграф {i} будет удалён: '{text}'")

        # Удаляем параграфы в обратном порядке (чтобы индексы не сбились)
        for i in reversed(paragraphs_to_remove):
            p = doc.paragraphs[i]._element
            p.getparent().remove(p)

        logger.info(f"[remove_empty_paragraphs] Удалено {len(paragraphs_to_remove)} из {total_paragraphs} параграфов из body")

        # Удаляем пустые параграфы из headers и footers всех секций
        headers_footers_removed = 0
        for section in doc.sections:
            # Обрабатываем header
            header_paras_to_remove = []
            for i, paragraph in enumerate(section.header.paragraphs):
                text = paragraph.text.strip()

                # Проверяем наличие маркера {# keep_empty #}
                if '{# keep_empty #}' in text:
                    _remove_marker_from_paragraph(paragraph)
                    logger.debug(f"[remove_empty_paragraphs] Header параграф {i} сохранён с маркером keep_empty")
                    continue

                if not text or all(char in ' -–—' for char in text):
                    header_paras_to_remove.append(i)
                    logger.debug(f"[remove_empty_paragraphs] Header параграф {i} будет удалён: '{text}'")

            for i in reversed(header_paras_to_remove):
                p = section.header.paragraphs[i]._element
                p.getparent().remove(p)
                headers_footers_removed += 1

            # Обрабатываем footer
            footer_paras_to_remove = []
            for i, paragraph in enumerate(section.footer.paragraphs):
                text = paragraph.text.strip()

                # Проверяем наличие маркера {# keep_empty #}
                if '{# keep_empty #}' in text:
                    _remove_marker_from_paragraph(paragraph)
                    logger.debug(f"[remove_empty_paragraphs] Footer параграф {i} сохранён с маркером keep_empty")
                    continue

                if not text or all(char in ' -–—' for char in text):
                    footer_paras_to_remove.append(i)
                    logger.debug(f"[remove_empty_paragraphs] Footer параграф {i} будет удалён: '{text}'")

            for i in reversed(footer_paras_to_remove):
                p = section.footer.paragraphs[i]._element
                p.getparent().remove(p)
                headers_footers_removed += 1

        if headers_footers_removed > 0:
            logger.info(f"[remove_empty_paragraphs] Удалено {headers_footers_removed} параграфов из headers/footers")

        # Сохраняем в BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    except Exception as e:
        logger.error(f"[remove_empty_paragraphs] Ошибка: {e}", exc_info=True)
        return doc_bytes


def remove_empty_table_rows(doc_bytes: bytes) -> bytes:
    """
    Удаляет пустые строки из таблиц в DOCX документе.

    Строка считается пустой, если все её ячейки пусты.

    Args:
        doc_bytes: DOCX документ в виде байтов

    Returns:
        Очищенный DOCX документ в виде байтов
    """
    try:
        doc = Document(BytesIO(doc_bytes))
        total_removed = 0

        for table_idx, table in enumerate(doc.tables):
            rows_to_remove = []

            for i, row in enumerate(table.rows):
                # Проверяем все ячейки в строке
                all_empty = True
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text and cell_text not in ['', '-', '—', '–']:
                        all_empty = False
                        break

                if all_empty:
                    rows_to_remove.append(i)

            # Удаляем строки в обратном порядке
            for i in reversed(rows_to_remove):
                table._tbl.remove(table.rows[i]._tr)
                total_removed += 1

            if rows_to_remove:
                logger.debug(f"[remove_empty_table_rows] Таблица {table_idx}: удалено {len(rows_to_remove)} строк")

        logger.info(f"[remove_empty_table_rows] Всего удалено {total_removed} пустых строк из таблиц")

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        return buffer.getvalue()

    except Exception as e:
        logger.error(f"[remove_empty_table_rows] Ошибка: {e}", exc_info=True)
        return doc_bytes


def clean_document(doc_bytes: bytes, remove_empty_rows: bool = True) -> bytes:
    """
    Полная очистка документа: удаление пустых параграфов и строк таблиц.

    Args:
        doc_bytes: DOCX документ в виде байтов
        remove_empty_rows: Удалять ли пустые строки таблиц (по умолчанию True)

    Returns:
        Очищенный DOCX документ в виде байтов
    """
    try:
        # Сначала удаляем пустые параграфы
        doc_bytes = remove_empty_paragraphs(doc_bytes)

        # Потом удаляем пустые строки таблиц (если нужно)
        if remove_empty_rows:
            doc_bytes = remove_empty_table_rows(doc_bytes)

        return doc_bytes

    except Exception as e:
        logger.error(f"[clean_document] Ошибка при очистке документа: {e}", exc_info=True)
        # В случае ошибки возвращаем оригинальный документ
        return doc_bytes
