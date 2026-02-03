# -*- coding: utf-8 -*-
"""
📄 Генераторы всех документов обучения на производстве

Содержит функции генерации для всех типов документов:
- Заявление (application)
- Приказ (order)
- Карточка теории (theory_card)
- Представление (presentation)
- Протокол комиссии (protocol)
- Заявление на пробную работу (trial_application)
- Заключение по пробной работе (trial_conclusion)
- Дневник обучения (diary)
"""
import logging
import re
from io import BytesIO
from typing import Dict, Any, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
from docxtpl import DocxTemplate
from django.utils import timezone

from .base import generate_training_document, get_template_path, _build_person_context, prepare_training_context
from directory.utils.commission_service import find_appropriate_commission
from directory.utils.declension import (
    decline_full_name,
    decline_phrase,
    get_initials_from_name,
    get_initials_before_surname,
)

logger = logging.getLogger(__name__)


def _capitalize_first(value: str) -> str:
    if not value:
        return ''
    return value[0].upper() + value[1:]


def _lowercase_first(value: str) -> str:
    if not value:
        return ''
    return value[0].lower() + value[1:]


# ============================================================================
# ЗАЯВЛЕНИЕ (application.docx)
# ============================================================================

def generate_application(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует заявление сотрудника на обучение.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('application.docx')

    context = dict(custom_context or {})
    base_context = prepare_training_context(training)
    director_ctx = dict(base_context.get('director') or {})
    if director_ctx:
        if not director_ctx.get('is_missing'):
            short_dative = get_initials_from_name(director_ctx.get('full_name_dative', ''))
            if short_dative:
                director_ctx['full_name_dative'] = short_dative
            position_dative = director_ctx.get('position_dative', '')
            if position_dative:
                director_ctx['position_dative'] = _capitalize_first(position_dative)
        if isinstance(context.get('director'), dict):
            merged = director_ctx.copy()
            merged.update(context['director'])
            director_ctx = merged
        context['director'] = director_ctx

    prior_qualification = getattr(training, 'prior_qualification', '') or ''
    if training.employee:
        if not prior_qualification:
            prior_qualification = training.employee.prior_qualification or ''
        if not prior_qualification:
            number = training.employee.qualification_document_number or ''
            date_value = training.employee.qualification_document_date
            date_str = date_value.strftime('%d.%m.%Y') if date_value else ''
            parts = []
            if number:
                parts.append(str(number))
            if date_str:
                parts.append(f'от {date_str}')
            prior_qualification = ' '.join(parts).strip()
        employee_ctx = dict(base_context.get('employee') or {})
        short_name = get_initials_before_surname(training.employee.full_name_nominative)
        if short_name:
            employee_ctx['short_name'] = short_name
        if isinstance(context.get('employee'), dict):
            merged = employee_ctx.copy()
            merged.update(context['employee'])
            employee_ctx = merged
        context['employee'] = employee_ctx
    context['prior_qualification'] = prior_qualification
    context['qualification_document_number'] = ''
    context['qualification_document_date'] = ''
    context['qualification_document_date_formatted'] = ''
    if training.start_date:
        context['application_date'] = training.start_date.strftime('%d.%m.%Y')

    return generate_training_document(
        training=training,
        template_path=template_path,
        document_name='Заявление',
        user=user,
        custom_context=context,
        use_vml=False
    )


# ============================================================================
# ПРИКАЗ НА ОБУЧЕНИЕ (order.docx)
# ============================================================================

def generate_order(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует приказ на обучение.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('order.docx')

    # Добавляем специфичный контекст для приказа
    context = dict(custom_context or {})
    base_context = prepare_training_context(training)
    if 'order_number' not in context and training.registration_number:
        context['order_number'] = training.registration_number
    if training.start_date:
        context['order_date'] = training.start_date.strftime('%d.%m.%Y')
        context['start_date'] = training.start_date.strftime('%d.%m.%Y')

    def adapt_role(role_key: str):
        role_ctx = dict(base_context.get(role_key) or {})
        if not role_ctx:
            return
        if role_ctx.get('is_missing'):
            context[role_key] = role_ctx
            return
        full_name = role_ctx.get('full_name_nominative', '')
        full_name_acc = role_ctx.get('full_name_accusative', '') or decline_full_name(full_name, 'accs')
        if full_name_acc:
            role_ctx['short_name'] = get_initials_from_name(full_name_acc)
            role_ctx['full_name_accusative'] = full_name_acc
        position_acc = role_ctx.get('position_accusative', '') or decline_phrase(
            role_ctx.get('position_nominative', ''),
            'accs',
        )
        if position_acc:
            role_ctx['position_accusative'] = position_acc
            role_ctx['position_nominative'] = _lowercase_first(position_acc)
        if isinstance(context.get(role_key), dict):
            merged = role_ctx.copy()
            merged.update(context[role_key])
            role_ctx = merged
        context[role_key] = role_ctx

    adapt_role('training_supervisor')
    adapt_role('instructor')
    adapt_role('theory_consultant')

    return generate_training_document(
        training=training,
        template_path=template_path,
        document_name='Приказ',
        user=user,
        custom_context=context,
        use_vml=False
    )


# ============================================================================
# КАРТОЧКА ТЕОРИИ (theory_card.docx)
# ============================================================================

def generate_theory_card(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует карточку теоретического обучения.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('theory_card.docx')

    context = prepare_training_context(training)
    if custom_context:
        context.update(custom_context)

    # Обработка разряда: склоняем слово в скобках в родительный падеж
    grade_ctx = context.get('qualification_grade')
    if grade_ctx and isinstance(grade_ctx, dict):
        label_ru = grade_ctx.get('label_ru', '') or ''
        if label_ru:
            match = re.search(r'(\d+)\s*\(([^)]+)\)', label_ru)
            if match:
                number = match.group(1)
                word = match.group(2)
                word_gen = decline_phrase(word, 'gent').lower()
                grade_ctx['label_ru'] = f"{number} ({word_gen})"

    template = DocxTemplate(template_path)
    template.render(context)
    rendered = BytesIO()
    template.save(rendered)
    rendered.seek(0)

    document = Document(rendered)
    sessions = context.get('consultation_sessions') or []

    def set_cell_text_with_font(cell, text, font_name='Times New Roman', font_size=14, align_center=False):
        cell.text = str(text or '')
        paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
        if align_center:
            paragraph.alignment = 1
        for run in paragraph.runs:
            run.font.name = font_name
            rpr = run._element.get_or_add_rPr()
            rfonts = rpr.get_or_add_rFonts()
            rfonts.set(qn('w:ascii'), font_name)
            rfonts.set(qn('w:hAnsi'), font_name)
            rfonts.set(qn('w:cs'), font_name)
            rfonts.set(qn('w:eastAsia'), font_name)
            run.font.size = Pt(font_size)

    def apply_table_borders(target_table):
        try:
            target_table.style = 'Table Grid'
            return
        except Exception:
            pass

        from docx.oxml import OxmlElement

        def set_cell_border(cell, **kwargs):
            tc = cell._tc
            tc_pr = tc.get_or_add_tcPr()
            tc_borders = tc_pr.first_child_found_in('w:tcBorders')
            if tc_borders is None:
                tc_borders = OxmlElement('w:tcBorders')
                tc_pr.append(tc_borders)
            for edge, edge_data in kwargs.items():
                edge_tag = qn(f'w:{edge}')
                element = tc_borders.find(edge_tag)
                if element is None:
                    element = OxmlElement(f'w:{edge}')
                    tc_borders.append(element)
                for key, val in edge_data.items():
                    element.set(qn(f'w:{key}'), str(val))

        border_spec = {
            'sz': '8',
            'val': 'single',
            'color': '000000',
        }
        for row in target_table.rows:
            for cell in row.cells:
                set_cell_border(
                    cell,
                    top=border_spec,
                    bottom=border_spec,
                    left=border_spec,
                    right=border_spec,
                )

    # Заполняем таблицу консультаций
    for table in document.tables:
        if not table.rows:
            continue
        header_text = ''.join(cell.text for cell in table.rows[0].cells)
        if 'Дата занятий' not in header_text:
            continue

        row_start = 2  # после заголовка и строки с номерами
        total_rows = row_start + len(sessions) + 1
        while len(table.rows) < total_rows:
            table.add_row()

        for idx, session in enumerate(sessions):
            row = table.rows[row_start + idx]
            set_cell_text_with_font(row.cells[0], session.get('date', ''))
            set_cell_text_with_font(row.cells[1], session.get('kind', 'Консультация'))
            set_cell_text_with_font(row.cells[2], session.get('hours', ''), align_center=True)
            set_cell_text_with_font(row.cells[3], session.get('consultant_initials', ''))
            set_cell_text_with_font(row.cells[4], '✓')

        total_row = table.rows[row_start + len(sessions)]
        set_cell_text_with_font(total_row.cells[1], 'Итого')
        set_cell_text_with_font(
            total_row.cells[2],
            context.get('total_consultation_hours', ''),
            align_center=True
        )

        apply_table_borders(table)
        break

    output = BytesIO()
    document.save(output)
    output.seek(0)

    employee_name = training.employee.full_name_nominative if training.employee else 'Без_сотрудника'
    safe_name = employee_name.replace(' ', '_')
    timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Карточка_теории_{safe_name}_{timestamp}.docx"

    return {
        'content': output,
        'filename': filename
    }


# ============================================================================
# ПРЕДСТАВЛЕНИЕ (presentation.docx)
# ============================================================================

def generate_presentation(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует представление на сотрудника.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('presentation.docx')

    base_context = prepare_training_context(training)
    context = dict(custom_context or {})

    # Обработка разряда: склоняем слово в скобках в родительный падеж
    # "3 (третий)" -> "третьего разряда"
    grade_ctx = dict(base_context.get('qualification_grade') or {})
    if grade_ctx:
        label_ru = grade_ctx.get('label_ru', '') or ''
        if label_ru:
            match = re.search(r'\(([^)]+)\)', label_ru)
            if match:
                word = match.group(1)
                word_gen = decline_phrase(word, 'gent').lower()
                grade_ctx['label_ru_genitive'] = f"{word_gen} разряда"
            else:
                grade_ctx['label_ru_genitive'] = decline_phrase(label_ru, 'gent').lower()
        if isinstance(context.get('qualification_grade'), dict):
            merged = grade_ctx.copy()
            merged.update(context['qualification_grade'])
            grade_ctx = merged
        context['qualification_grade'] = grade_ctx

    return generate_training_document(
        training=training,
        template_path=template_path,
        document_name='Представление',
        user=user,
        custom_context=context,
        use_vml=False
    )


# ============================================================================
# ПРОТОКОЛ КОМИССИИ (protocol.docx)
# ============================================================================

def generate_protocol(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует протокол квалификационной комиссии.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('protocol.docx')

    # Добавляем специфичный контекст для протокола
    context = custom_context or {}

    # Номер и дата протокола
    if 'protocol_number' not in context and training.protocol_number:
        context['protocol_number'] = training.protocol_number
    if 'protocol_date' not in context and training.protocol_date:
        context['protocol_date'] = training.protocol_date.strftime('%d.%m.%Y')

    # Состав комиссии (квалификационная комиссия)
    commission = training.commission
    if not commission and training.employee:
        try:
            commission = find_appropriate_commission(
                training.employee,
                commission_type='qualification'
            )
        except Exception:
            commission = None

    if commission:
        members = commission.members.select_related('employee', 'employee__position').filter(is_active=True)

        # Председатель
        chairman = members.filter(role='chairman').first()
        if not chairman:
            chairman = members.filter(role='vice_chairman').first()

        if chairman and chairman.employee:
            context['commission_chairman_name'] = chairman.employee.full_name_nominative
            if chairman.employee.position:
                context['commission_chairman_position'] = chairman.employee.position.position_name
            context['commission_chairman'] = _build_person_context(chairman.employee)
        else:
            # Явно очищаем, чтобы не подтягивались значения из training.commission_chairman.
            context['commission_chairman'] = _build_person_context(None)

        # Члены комиссии (включая заместителя, если он не председатель)
        regular_members = members.filter(role__in=['member', 'vice_chairman'])
        context['commission_members_list'] = [
            {
                'name': m.employee.full_name_nominative,
                'position': m.employee.position.position_name if m.employee.position else ''
            }
            for m in regular_members
            if m.employee_id
        ]
        context['commission_members'] = [
            _build_person_context(m.employee)
            for m in regular_members
            if m.employee_id
        ]

    return generate_training_document(
        training=training,
        template_path=template_path,
        document_name='Протокол',
        user=user,
        custom_context=context,
        use_vml=False
    )


# ============================================================================
# ЗАЯВЛЕНИЕ НА ПРОБНУЮ РАБОТУ (trial_application.docx)
# ============================================================================

def generate_trial_application(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует заявление на допуск к пробной работе.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('trial_application.docx')

    base_context = prepare_training_context(training)
    # Добавляем дату пробной работы
    context = dict(custom_context or {})

    org_ctx = dict(base_context.get('organization') or {})
    if isinstance(context.get('organization'), dict):
        merged = org_ctx.copy()
        merged.update(context['organization'])
        org_ctx = merged
    if org_ctx:
        context['organization'] = org_ctx
        if org_ctx.get('short_name_ru') and 'organization_short_name_ru' not in context:
            context['organization_short_name_ru'] = org_ctx['short_name_ru']

    chairman_ctx = dict(base_context.get('commission_chairman') or {})
    if chairman_ctx and not chairman_ctx.get('is_missing'):
        full_name_dat = chairman_ctx.get('full_name_dative', '')
        if full_name_dat:
            chairman_ctx['full_name_dative'] = get_initials_from_name(full_name_dat)
    if isinstance(context.get('commission_chairman'), dict):
        merged = chairman_ctx.copy()
        merged.update(context['commission_chairman'])
        chairman_ctx = merged
    if chairman_ctx:
        context['commission_chairman'] = chairman_ctx

    grade_ctx = dict(base_context.get('qualification_grade') or {})
    if grade_ctx:
        label_ru = grade_ctx.get('label_ru', '') or ''
        if label_ru.lower().startswith('по '):
            grade_phrase = label_ru
        else:
            ordinal_map = {
                1: 'первый',
                2: 'второй',
                3: 'третий',
                4: 'четвертый',
                5: 'пятый',
                6: 'шестой',
            }
            grade_number = grade_ctx.get('grade_number')
            ordinal = ordinal_map.get(grade_number, '')
            if not ordinal and label_ru:
                match = re.search(r'\(([^)]+)\)', label_ru)
                ordinal = match.group(1).strip() if match else label_ru.strip()
            phrase = ordinal or ''
            if phrase and 'разряд' not in phrase:
                phrase = f"{phrase} разряд"
            grade_phrase = f"по {decline_phrase(phrase, 'datv')}" if phrase else ''
        if grade_phrase:
            grade_ctx['label_ru'] = grade_phrase
    if isinstance(context.get('qualification_grade'), dict):
        merged = grade_ctx.copy()
        merged.update(context['qualification_grade'])
        grade_ctx = merged
    if grade_ctx:
        context['qualification_grade'] = grade_ctx
        if grade_ctx.get('label_ru') and 'qualification_grade_ru' not in context:
            context['qualification_grade_ru'] = grade_ctx['label_ru']

    if 'practical_date' not in context and training.practical_date:
        context['practical_date'] = training.practical_date.strftime('%d.%m.%Y')
    if 'practical_work_topic' not in context and training.practical_work_topic:
        context['practical_work_topic'] = training.practical_work_topic

    return generate_training_document(
        training=training,
        template_path=template_path,
        document_name='Заявление_на_пробную_работу',
        user=user,
        custom_context=context,
        use_vml=False
    )


# ============================================================================
# ЗАКЛЮЧЕНИЕ ПО ПРОБНОЙ РАБОТЕ (trial_conclusion.docx)
# ============================================================================

def generate_trial_conclusion(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует заключение по пробной работе.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    template_path = get_template_path('trial_conclusion.docx')

    base_context = prepare_training_context(training)
    # Добавляем результаты пробной работы
    context = dict(custom_context or {})

    profession_ctx = dict(base_context.get('profession') or {})
    if profession_ctx:
        gen_name = profession_ctx.get('name_ru_genitive') or profession_ctx.get('name_ru_nominative', '')
        if gen_name:
            profession_ctx['name_ru_nominative'] = _lowercase_first(gen_name)
        if isinstance(context.get('profession'), dict):
            merged = profession_ctx.copy()
            merged.update(context['profession'])
            profession_ctx = merged
        context['profession'] = profession_ctx
        if profession_ctx.get('name_ru_nominative') and 'profession_nominative_ru' not in context:
            context['profession_nominative_ru'] = profession_ctx['name_ru_nominative']

    grade_ctx = dict(base_context.get('qualification_grade') or {})
    if grade_ctx:
        label_ru = grade_ctx.get('label_ru', '') or ''
        grade_label = label_ru
        if label_ru:
            match = re.search(r'(\d+)\s*\(([^)]+)\)', label_ru)
            if match:
                number = match.group(1)
                word = match.group(2)
                word_gen = decline_phrase(word, 'gent').lower()
                grade_label = f"{number} ({word_gen})"
            else:
                grade_label = decline_phrase(label_ru, 'gent').lower()
        grade_ctx['label_ru'] = grade_label
        if isinstance(context.get('qualification_grade'), dict):
            merged = grade_ctx.copy()
            merged.update(context['qualification_grade'])
            grade_ctx = merged
        context['qualification_grade'] = grade_ctx
        if grade_ctx.get('label_ru') and 'qualification_grade_ru' not in context:
            context['qualification_grade_ru'] = grade_ctx['label_ru']

    if 'practical_score' not in context and training.practical_score:
        context['practical_score'] = training.practical_score
    if 'practical_work_topic' not in context and training.practical_work_topic:
        context['practical_work_topic'] = training.practical_work_topic

    return generate_training_document(
        training=training,
        template_path=template_path,
        document_name='Заключение_пробная_работа',
        user=user,
        custom_context=context,
        use_vml=False
    )


# ============================================================================
# ДНЕВНИК ОБУЧЕНИЯ (diary_podgotovka/diary_perepodgotovka)
# ============================================================================

def generate_diary(training, user=None, custom_context: Optional[Dict[str, Any]] = None):
    """
    Генерирует дневник обучения (подготовка или переподготовка).

    Выбор шаблона происходит автоматически на основе типа обучения:
    - preparation → diary_podgotovka_voditel_pogruzchika.docx
    - retraining → diary_perepodgotovka_voditel_pogruzchika.docx

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)
        custom_context: Дополнительный контекст (опционально)

    Returns:
        Dict или None
    """
    # Получаем путь к шаблону дневника из модели
    template_path = training.get_diary_template_path()

    if not template_path:
        logger.error(f"Не найден шаблон дневника для обучения {training.id}")
        return None

    template_doc = Document(template_path)
    use_name_placeholder = False
    use_period_placeholder = False
    if template_doc.tables:
        if template_doc.tables[0].rows:
            use_name_placeholder = '{{' in template_doc.tables[0].rows[0].cells[0].text
        if len(template_doc.tables) > 1 and template_doc.tables[1].rows:
            use_period_placeholder = '{{' in template_doc.tables[1].rows[0].cells[0].text

    context = prepare_training_context(training)
    if custom_context:
        context.update(custom_context)
    diary_entries = context.get('diary_entries') or (
        training.get_diary_entries() if hasattr(training, 'get_diary_entries') else []
    )
    context['diary_entries'] = diary_entries

    # Рендерим Jinja-плейсхолдеры, затем дополняем таблицу дат
    template = DocxTemplate(template_path)
    template.render(context)
    rendered = BytesIO()
    template.save(rendered)
    rendered.seek(0)

    document = Document(rendered)

    def replace_placeholder(paragraph, pattern, value):
        """Возвращает True если замена произошла."""
        if not paragraph.runs:
            return False
        full_text = ''.join(run.text for run in paragraph.runs)
        if not re.search(pattern, full_text):
            return False
        new_text = re.sub(pattern, value, full_text)
        paragraph.runs[0].text = new_text
        for extra in paragraph.runs[1:]:
            extra.text = ''
        return True

    def add_section_break_after_paragraph(paragraph):
        """Добавляет разрыв раздела (новая страница) после параграфа."""
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn as docx_qn
        from docx.enum.section import WD_ORIENT

        # Создаём sectPr с типом nextPage
        sectPr = OxmlElement('w:sectPr')
        sect_type = OxmlElement('w:type')
        sect_type.set(docx_qn('w:val'), 'nextPage')
        sectPr.append(sect_type)

        # Добавляем в конец pPr параграфа
        pPr = paragraph._element.get_or_add_pPr()
        pPr.append(sectPr)

    instructor_short_name = context.get('instructor', {}).get('short_name', '')
    placeholder_pattern = r'{{\s*instructor\.short_name\s*}}'
    instructor_paragraph_found = None

    for paragraph in document.paragraphs:
        if replace_placeholder(paragraph, placeholder_pattern, instructor_short_name):
            instructor_paragraph_found = paragraph
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_placeholder(paragraph, placeholder_pattern, instructor_short_name):
                        instructor_paragraph_found = paragraph

    # Добавляем разрыв раздела после параграфа с instructor.short_name
    if instructor_paragraph_found:
        add_section_break_after_paragraph(instructor_paragraph_found)

    if document.tables:
        def set_cell_text_preserve(cell, text, font_name=None, font_size=None):
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            if paragraph.runs:
                paragraph.runs[0].text = text
                for extra in paragraph.runs[1:]:
                    extra.text = ''
            else:
                paragraph.add_run(text)
            if font_name or font_size:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(text)
                if font_name:
                    run.font.name = font_name
                    rpr = run._element.get_or_add_rPr()
                    rfonts = rpr.get_or_add_rFonts()
                    rfonts.set(qn('w:ascii'), font_name)
                    rfonts.set(qn('w:hAnsi'), font_name)
                    rfonts.set(qn('w:cs'), font_name)
                    rfonts.set(qn('w:eastAsia'), font_name)
                if font_size:
                    run.font.size = Pt(font_size)

        def set_cell_text_with_font(cell, text, font_name, font_size):
            cell.text = text
            paragraph = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
            for run in paragraph.runs:
                run.font.name = font_name
                rpr = run._element.get_or_add_rPr()
                rfonts = rpr.get_or_add_rFonts()
                rfonts.set(qn('w:ascii'), font_name)
                rfonts.set(qn('w:hAnsi'), font_name)
                rfonts.set(qn('w:cs'), font_name)
                rfonts.set(qn('w:eastAsia'), font_name)
                run.font.size = Pt(font_size)

        # ФИО сотрудника
        if training.employee and document.tables[0].rows and not use_name_placeholder:
            set_cell_text_preserve(
                document.tables[0].rows[0].cells[0],
                training.employee.full_name_nominative,
            )

        # Период обучения
        start_date = getattr(training, 'start_date', None)
        end_date = getattr(training, 'end_date', None)
        if start_date and end_date and len(document.tables) > 1 and not use_period_placeholder:
            period = f"c {start_date.strftime('%d.%m.%Y')} по {end_date.strftime('%d.%m.%Y')}"
            set_cell_text_preserve(document.tables[1].rows[0].cells[0], period)

        # Даты в дневнике
        if len(document.tables) > 2:
            table = document.tables[2]
            row_start = 2  # первые две строки — заголовки
            for idx, entry in enumerate(diary_entries):
                row_idx = row_start + idx
                if row_idx >= len(table.rows):
                    break
                date_value = entry.get('date')
                if date_value:
                    set_cell_text_with_font(
                        table.rows[row_idx].cells[1],
                        date_value.strftime('%d.%m.%Y'),
                        font_name='Times New Roman',
                        font_size=12,
                    )
                    # Галочка в 7-м столбце (индекс 6)
                    if len(table.rows[row_idx].cells) > 6:
                        set_cell_text_with_font(
                            table.rows[row_idx].cells[6],
                            '✓',
                            font_name='Times New Roman',
                            font_size=12,
                        )

    output = BytesIO()
    document.save(output)
    output.seek(0)

    employee_name = training.employee.full_name_nominative if training.employee else 'Без_сотрудника'
    safe_name = employee_name.replace(' ', '_')
    timestamp = timezone.now().strftime('%Y%m%d_%H%M%S')
    filename = f"Дневник_{safe_name}_{timestamp}.docx"

    return {
        'content': output,
        'filename': filename
    }


# ============================================================================
# ГЕНЕРАЦИЯ ВСЕХ ДОКУМЕНТОВ
# ============================================================================

# Маппинг типов документов на нумерованные префиксы (соответствует шаблонам)
DOCUMENT_PREFIXES = {
    'application': '01. Заявление',
    'order': '02. Приказ о назначении обучения',
    'theory_card': '03. Карточка теория',
    'diary': '04. Дневник',
    'trial_conclusion': '05. Заключение на пробную работу',
    'presentation': '06. Представление на квалификационную работу',
    'trial_application': '07. Заявление на квалификационный экзамен',
    'protocol': '08. Протокол квалификационной комиссии',
}


def generate_all_training_documents(training, user=None):
    """
    Генерирует все документы для обучения.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)

    Returns:
        Dict[str, Dict]: Словарь с результатами генерации
                         {'application': {...}, 'order': {...}, ...}
    """
    results = {}

    # Упорядоченный список генераторов (порядок соответствует нумерации)
    generators = [
        ('application', generate_application),
        ('order', generate_order),
        ('theory_card', generate_theory_card),
        ('diary', generate_diary),
        ('trial_application', generate_trial_application),
        ('trial_conclusion', generate_trial_conclusion),
        ('presentation', generate_presentation),
        ('protocol', generate_protocol),
    ]

    for doc_type, generator_func in generators:
        try:
            result = generator_func(training, user)
            if result:
                # Добавляем нумерованный префикс к имени файла
                prefix = DOCUMENT_PREFIXES.get(doc_type, '')
                if prefix and result.get('filename'):
                    result['filename'] = f"{prefix} - {result['filename']}"
                results[doc_type] = result
                logger.info(f"✅ {doc_type}: {result['filename']}")
            else:
                results[doc_type] = None
                logger.warning(f"❌ {doc_type}: не сгенерирован")
        except Exception as e:
            logger.error(f"❌ {doc_type}: ошибка - {e}")
            results[doc_type] = None

    return results


def generate_merged_document(training, user=None):
    """
    Генерирует все документы и объединяет их в один DOCX с разрывами страниц.

    Args:
        training: Объект TrainingAssignment
        user: Пользователь (опционально)

    Returns:
        Dict или None: {'content': BytesIO, 'filename': str}
    """
    from docxcompose.composer import Composer
    from docx import Document as DocxDocument
    from docx.shared import Cm
    from docx.oxml.ns import qn as docx_qn
    from lxml import etree

    # Получаем все документы
    results = generate_all_training_documents(training, user)

    # Фильтруем успешно сгенерированные и сортируем по номеру из DOCUMENT_PREFIXES
    valid_docs = []
    for doc_type, result in results.items():
        if result and result.get('content'):
            # Извлекаем номер из префикса (например, "01" из "01. Заявление")
            prefix = DOCUMENT_PREFIXES.get(doc_type, '99')
            try:
                sort_key = int(prefix.split('.')[0])
            except (ValueError, IndexError):
                sort_key = 99
            valid_docs.append((sort_key, doc_type, result))

    # Сортируем по номеру
    valid_docs.sort(key=lambda x: x[0])
    # Убираем sort_key из кортежа
    valid_docs = [(doc_type, result) for _, doc_type, result in valid_docs]

    if not valid_docs:
        logger.error("Нет документов для объединения")
        return None

    try:
        # Загружаем первый документ как базовый
        first_content = valid_docs[0][1]['content']
        first_content.seek(0)
        master = DocxDocument(first_content)

        # Настраиваем поля: левое 2 см, остальные 1 см
        for section in master.sections:
            section.left_margin = Cm(2)
            section.right_margin = Cm(1)
            section.top_margin = Cm(1)
            section.bottom_margin = Cm(1)

        composer = Composer(master)

        # Добавляем остальные документы
        for i, (doc_type, doc_result) in enumerate(valid_docs[1:], start=1):
            doc_content = doc_result['content']
            doc_content.seek(0)
            sub_doc = DocxDocument(doc_content)

            # Настраиваем поля для добавляемого документа
            for section in sub_doc.sections:
                section.left_margin = Cm(2)
                section.right_margin = Cm(1)
                section.top_margin = Cm(1)
                section.bottom_margin = Cm(1)

            # Добавляем разрыв страницы В НАЧАЛО документа (перед первым элементом)
            # Создаём параграф с разрывом страницы
            from docx.oxml.ns import qn as docx_qn
            from docx.oxml import OxmlElement

            # Создаём новый параграф с page break
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            br = OxmlElement('w:br')
            br.set(docx_qn('w:type'), 'page')
            r.append(br)
            p.append(r)

            # Вставляем в начало body документа
            body = sub_doc.element.body
            if len(body) > 0:
                body.insert(0, p)
            else:
                body.append(p)

            composer.append(sub_doc)

        # Сохраняем результат
        output = BytesIO()
        composer.save(output)
        output.seek(0)

        # Формируем имя файла
        employee_name = ''
        if training.employee:
            employee_name = training.employee.full_name_nominative.replace(' ', '_')
        filename = f"Документы_обучения_{employee_name}.docx"

        logger.info(f"✅ Объединено {len(valid_docs)} документов в один файл")
        return {'content': output, 'filename': filename}

    except Exception as e:
        logger.error(f"Ошибка объединения документов: {e}")
        import traceback
        logger.error(traceback.format_exc())
        return None
